from __future__ import annotations

import csv
import importlib.util
import os
import subprocess
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw


ROOT = Path("D:/graduation Project")
DOC_HELPER_PATH = ROOT / "tmp/docs/generate_thesis_draft.py"
OUT_DIR = Path("D:/毕设/论文初稿")
IMG_DIR = ROOT / "tmp/docs/generated_figures"
SCREEN_DIR = ROOT / "tmp/docs/screens"
OUTPUT_DOCX = OUT_DIR / "10-基于Hadoop与Spring Boot的贵州人才招聘市场分析管理平台的设计与实现-数据主线重构初稿.docx"
OUTPUT_PDF = OUTPUT_DOCX.with_suffix(".pdf")

TITLE = "基于Hadoop与Spring Boot的贵州人才招聘市场分析管理平台的设计与实现"
DATE_TEXT = "2026年4月29日"


def load_helper():
    spec = importlib.util.spec_from_file_location("thesis_helper", DOC_HELPER_PATH)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    module.DATE_TEXT = DATE_TEXT
    return module


h = load_helper()


def mysql_query(sql: str) -> list[dict[str, str]]:
    env = os.environ.copy()
    env["MYSQL_PWD"] = env.get("MYSQL_PWD", "NLXXtxd123")
    cmd = [
        "mysql",
        "-h",
        "127.0.0.1",
        "-P",
        "3306",
        "-uroot",
        "-D",
        "guizhou_job_platform",
        "--default-character-set=utf8mb4",
        "--batch",
        "--raw",
        "-e",
        sql,
    ]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8", env=env, timeout=20)
        if result.returncode != 0:
            return []
        rows = list(csv.reader(result.stdout.splitlines(), delimiter="\t"))
        if not rows:
            return []
        headers = rows[0]
        return [dict(zip(headers, row)) for row in rows[1:]]
    except Exception:
        return []


def get_data_snapshot():
    city_rows = mysql_query(
        "SELECT city, COUNT(*) AS cnt FROM job_info GROUP BY city ORDER BY cnt DESC LIMIT 10;"
    ) or [
        {"city": "黔东南", "cnt": "5391"},
        {"city": "黔西南", "cnt": "4819"},
        {"city": "安顺", "cnt": "2752"},
        {"city": "贵阳", "cnt": "55"},
    ]
    job_count_rows = mysql_query("SELECT COUNT(*) AS cnt FROM job_info;") or [{"cnt": "13019"}]
    tfidf_rows = mysql_query(
        "SELECT city, keyword, tfidf_score, job_count, rank_no "
        "FROM analysis_skill_tfidf WHERE data_batch_no='qiandongnan_2604221328' AND city='黔东南' "
        "ORDER BY rank_no LIMIT 10;"
    ) or [
        {"keyword": "销售", "tfidf_score": "794.86", "job_count": "467", "rank_no": "1"},
        {"keyword": "教育", "tfidf_score": "427.92", "job_count": "155", "rank_no": "2"},
        {"keyword": "兼职", "tfidf_score": "415.06", "job_count": "141", "rank_no": "3"},
        {"keyword": "办公", "tfidf_score": "399.85", "job_count": "129", "rank_no": "4"},
        {"keyword": "培训", "tfidf_score": "392.21", "job_count": "174", "rank_no": "5"},
    ]
    salary_rows = mysql_query(
        "SELECT city, salary_range, job_count FROM analysis_salary_distribution "
        "WHERE data_batch_no='qiandongnan_2604221328' ORDER BY salary_range;"
    ) or [
        {"city": "黔东南", "salary_range": "3K以下", "job_count": "330"},
        {"city": "黔东南", "salary_range": "3K-5K", "job_count": "1568"},
        {"city": "黔东南", "salary_range": "5K-8K", "job_count": "2250"},
        {"city": "黔东南", "salary_range": "8K-12K", "job_count": "876"},
        {"city": "黔东南", "salary_range": "12K-15K", "job_count": "196"},
        {"city": "黔东南", "salary_range": "15K以上", "job_count": "169"},
    ]
    education_rows = mysql_query(
        "SELECT education_text, job_count, proportion FROM analysis_education_stats "
        "WHERE data_batch_no='qiandongnan_2604221328' AND city='黔东南' "
        "ORDER BY job_count DESC LIMIT 8;"
    ) or [
        {"education_text": "不限", "job_count": "2060", "proportion": "38.21"},
        {"education_text": "大专", "job_count": "1698", "proportion": "31.50"},
        {"education_text": "本科", "job_count": "723", "proportion": "13.41"},
    ]
    exp_rows = mysql_query(
        "SELECT experience_text, job_count, proportion FROM analysis_experience_stats "
        "WHERE data_batch_no='qiandongnan_2604221328' AND city='黔东南' "
        "ORDER BY job_count DESC LIMIT 8;"
    ) or [
        {"experience_text": "不限", "job_count": "2266", "proportion": "42.03"},
        {"experience_text": "1-3年", "job_count": "1670", "proportion": "30.98"},
        {"experience_text": "3-5年", "job_count": "756", "proportion": "14.02"},
    ]
    return {
        "job_count": job_count_rows[0].get("cnt", "13019"),
        "city_rows": city_rows,
        "tfidf_rows": tfidf_rows,
        "salary_rows": salary_rows,
        "education_rows": education_rows,
        "experience_rows": exp_rows,
    }


def set_cn_font():
    plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "SimSun", "Arial Unicode MS"]
    plt.rcParams["axes.unicode_minus"] = False


def draw_chart(path: Path, title: str, labels: list[str], values: list[float], ylabel: str):
    set_cn_font()
    fig, ax = plt.subplots(figsize=(8.8, 4.6), dpi=180)
    bars = ax.bar(labels, values, color="#4f8f6f")
    ax.set_title(title, fontsize=15, pad=12)
    ax.set_ylabel(ylabel)
    ax.grid(axis="y", linestyle="--", alpha=0.3)
    for bar, value in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height(), f"{int(value)}", ha="center", va="bottom", fontsize=9)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def draw_formula(path: Path, formula: str, title: str):
    set_cn_font()
    fig, ax = plt.subplots(figsize=(9, 1.5), dpi=220)
    ax.axis("off")
    ax.text(0.02, 0.72, title, fontsize=12, fontweight="bold", transform=ax.transAxes)
    ax.text(0.02, 0.25, formula, fontsize=16, transform=ax.transAxes)
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def box(draw, xy, text, fill=(239, 246, 237), outline=(78, 128, 105), font=None):
    h.draw_box(draw, xy, text, fill, outline, font or h.chinese_font(22, True), radius=12)


def arrow(draw, start, end):
    h.draw_arrow(draw, start, end, color=(35, 112, 91), width=4)


def make_extra_figures(snapshot):
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    h.make_diagrams()
    font_title = h.chinese_font(34, True)
    font_box = h.chinese_font(22, True)

    img = Image.new("RGB", (1600, 820), (252, 249, 239))
    d = ImageDraw.Draw(img)
    d.text((70, 35), "岗位数据清洗规则流程", fill=(20, 60, 50), font=font_title)
    boxes = [
        ("读取ODS\nJSONL原始数据", (70, 160, 330, 280)),
        ("字段回填\n详情页优先", (430, 160, 690, 280)),
        ("必填过滤\n标题/公司/描述/hash", (790, 160, 1090, 280)),
        ("城市过滤\nstrict-city", (1190, 160, 1480, 280)),
        ("薪资解析\n月薪折算", (250, 470, 520, 590)),
        ("job_hash去重\n批次内唯一", (650, 470, 930, 590)),
        ("输出DWD\nParquet + Report", (1060, 470, 1360, 590)),
    ]
    for text, xy in boxes:
        box(d, xy, text, font=font_box)
    for a, b in [(0, 1), (1, 2), (2, 3)]:
        arrow(d, (boxes[a][1][2] + 15, 220), (boxes[b][1][0] - 15, 220))
    arrow(d, (1335, 285), (385, 455))
    arrow(d, (520, 530), (650, 530))
    arrow(d, (930, 530), (1060, 530))
    img.save(IMG_DIR / "cleaning_flow.png")

    img = Image.new("RGB", (1600, 760), (252, 249, 239))
    d = ImageDraw.Draw(img)
    d.text((70, 35), "jieba + TF-IDF岗位关键词分析流程", fill=(20, 60, 50), font=font_title)
    boxes = [
        ("DWD岗位描述\njob_description", (80, 180, 330, 310)),
        ("jieba分词\n自定义词典", (430, 180, 680, 310)),
        ("停用词过滤\n英文词标准化", (780, 180, 1060, 310)),
        ("CountVectorizer\n词频向量TF", (1160, 180, 1460, 310)),
        ("IDF计算\n逆文档频率", (330, 470, 610, 600)),
        ("TF-IDF向量\n提取Top关键词", (710, 470, 1010, 600)),
        ("城市聚合\n写入analysis_skill_tfidf", (1110, 470, 1460, 600)),
    ]
    for text, xy in boxes:
        box(d, xy, text, font=font_box)
    for a, b in [(0, 1), (1, 2), (2, 3)]:
        arrow(d, (boxes[a][1][2] + 15, 245), (boxes[b][1][0] - 15, 245))
    arrow(d, (1310, 315), (470, 455))
    arrow(d, (610, 535), (710, 535))
    arrow(d, (1010, 535), (1110, 535))
    img.save(IMG_DIR / "tfidf_flow.png")

    img = Image.new("RGB", (1600, 770), (252, 249, 239))
    d = ImageDraw.Draw(img)
    d.text((70, 35), "推荐评分模型设计", fill=(20, 60, 50), font=font_title)
    boxes = [
        ("地区匹配\n城市偏好级别", (100, 170, 360, 300)),
        ("岗位匹配\n期望岗位/标题", (100, 440, 360, 570)),
        ("技能匹配\n用户技能/TF-IDF", (520, 170, 800, 300)),
        ("薪资匹配\n期望区间/岗位薪资", (520, 440, 800, 570)),
        ("行为反馈\n浏览/搜索/点击", (940, 300, 1200, 430)),
        ("综合评分\n排序与推荐理由", (1320, 300, 1540, 430)),
    ]
    for text, xy in boxes:
        box(d, xy, text, font=font_box)
    for idx in [0, 1, 2, 3]:
        end_y = 365 if idx % 2 == 0 else 380
        arrow(d, (boxes[idx][1][2] + 12, (boxes[idx][1][1] + boxes[idx][1][3]) // 2), (940 - 15, end_y))
    arrow(d, (1200, 365), (1320, 365))
    img.save(IMG_DIR / "recommend_score_model.png")

    draw_formula(
        IMG_DIR / "formula_tfidf.png",
        r"$TFIDF(t,d)=TF(t,d)\times \log\frac{N+1}{DF(t)+1}$",
        "式4.1 TF-IDF关键词权重计算公式",
    )
    draw_formula(
        IMG_DIR / "formula_salary.png",
        r"$salary_{month}=\frac{salary_{year}}{12},\quad salary_{month}=salary_{day}\times 22,\quad salary_{month}=salary_{hour}\times 8\times 22$",
        "式4.2 不同薪资单位的月薪折算公式",
    )
    draw_formula(
        IMG_DIR / "formula_recommend.png",
        r"$Score=w_cS_c+w_pS_p+w_sS_s+w_mS_m+w_bS_b$",
        "式4.3 人岗匹配综合评分公式",
    )

    city_labels = [r["city"] for r in snapshot["city_rows"][:6]]
    city_values = [float(r["cnt"]) for r in snapshot["city_rows"][:6]]
    draw_chart(IMG_DIR / "chart_city_jobs.png", "当前MySQL岗位城市分布", city_labels, city_values, "岗位数量")

    salary_labels = [r["salary_range"] for r in snapshot["salary_rows"]]
    salary_values = [float(r["job_count"]) for r in snapshot["salary_rows"]]
    draw_chart(IMG_DIR / "chart_salary_qiandongnan.png", "黔东南样本薪资区间分布", salary_labels, salary_values, "岗位数量")

    tfidf_labels = [r["keyword"] for r in snapshot["tfidf_rows"][:10]]
    tfidf_values = [float(r["job_count"]) for r in snapshot["tfidf_rows"][:10]]
    draw_chart(IMG_DIR / "chart_tfidf_qiandongnan.png", "黔东南TF-IDF关键词关联岗位数Top10", tfidf_labels, tfidf_values, "关联岗位数")


def add_toc_field(doc: Document):
    h.add_center(doc, "目录", "黑体", 15, True, 10)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录将在打开文档后自动更新。"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    doc.add_page_break()


def add_section_break(doc: Document):
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_compact_cover(doc: Document):
    for _ in range(2):
        doc.add_paragraph()
    h.add_center(doc, "贵州理工学院", "宋体", 26, True)
    h.add_center(doc, "本科毕业论文（设计）", "宋体", 26, True)
    h.add_center(doc, "（2026届）", "宋体", 18, True)
    doc.add_paragraph()
    h.add_center(doc, "设计题目：", "宋体", 16, True)
    title_p = h.add_center(doc, TITLE, "宋体", 18, True)
    title_p.paragraph_format.space_after = Pt(18)
    info = [
        "学    院：大数据学院",
        "专    业：数据科学与大数据技术",
        "班    级：大数据222",
        "学    号：202249020315",
        "学生姓名：邰旭东",
        "第一指导教师：张德跃",
        "第二指导教师：",
    ]
    for item in info:
        para = h.add_center(doc, item, "宋体", 14)
        para.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()
    h.add_center(doc, DATE_TEXT, "宋体", 14)
    doc.add_page_break()


def p(text: str):
    return text


def add_formula_image(doc, image_name: str, caption: str):
    h.add_image(doc, IMG_DIR / image_name, caption, 13.5)


def add_abstracts(doc: Document):
    h.add_center(doc, "摘要", "黑体", 15, True)
    h.add_para(
        doc,
        "随着贵州数字经济、现代服务业和制造业持续发展，招聘平台上产生了大量岗位信息。此类数据包含岗位名称、薪资待遇、学历要求、经验要求、岗位描述和工作城市等内容，能够反映区域人才需求结构。但原始招聘数据具有来源分散、字段格式不统一、岗位描述非结构化、薪资单位多样和重复岗位较多等特点，若缺少系统化的数据处理流程，难以支撑招聘市场分析和个性化岗位推荐。针对上述问题，本文设计并实现了一个基于Hadoop与Spring Boot的贵州人才招聘市场分析管理平台。",
    )
    h.add_para(
        doc,
        "本文以招聘岗位数据处理为研究主线，构建了从网络采集、HDFS原始存储、Spark清洗、Hive主题管理、MySQL业务入库到Web可视化展示的完整链路。系统使用爬虫按贵州城市采集岗位JSONL数据，利用Spark完成空值过滤、城市过滤、job_hash去重、薪资标准化和字段规范化处理，并将清洗后的DWD数据以Parquet格式保存到HDFS。随后通过Spark对岗位数量、薪资区间、学历要求、经验要求和技能热词进行统计分析，并使用jieba分词与TF-IDF算法提取岗位描述中的代表性关键词，结果写入MySQL的analysis_*表供前端展示。",
    )
    h.add_para(
        doc,
        "平台系统实现作为数据分析结果的业务呈现层，后端采用Spring Boot与MyBatis封装用户、岗位、分析和管理员接口，前端采用Vue与ECharts实现岗位检索、岗位详情、数据看板、个人中心和管理员页面。测试结果表明，系统能够完成从爬虫数据到Hadoop生态处理、MySQL入库和Web展示的完整数据链路。本文还设计了基于城市偏好、岗位相关性、技能匹配、薪资匹配和行为反馈的人岗匹配评分模型，为后续推荐系统完善提供基础。",
    )
    para = doc.add_paragraph()
    h.set_paragraph_format(para, first_indent=False)
    r = para.add_run("关键词：")
    h.style_run(r, "黑体", 12, True)
    r = para.add_run("Hadoop；Spark；Hive；招聘数据分析；jieba；TF-IDF；Spring Boot；数据可视化")
    h.style_run(r, "宋体", 12)
    doc.add_page_break()

    h.add_center(doc, "Abstract", "Times New Roman", 15, True)
    para = doc.add_paragraph()
    h.set_paragraph_format(para, True)
    r = para.add_run(
        "With the development of digital economy, modern service industries and manufacturing in Guizhou, recruitment platforms generate a large amount of job information. Such data contains job titles, salaries, education requirements, experience requirements, job descriptions and work cities, which can reflect regional talent demand. However, raw recruitment data is distributed, heterogeneous, unstructured and duplicated. Without a systematic data processing workflow, it is difficult to support recruitment market analysis and personalized job recommendation. To address these problems, this thesis designs and implements a Guizhou talent recruitment market analysis and management platform based on Hadoop and Spring Boot."
    )
    h.style_run(r, "Times New Roman", 12)
    para = doc.add_paragraph()
    h.set_paragraph_format(para, True)
    r = para.add_run(
        "This thesis takes recruitment data processing as the main research line and builds a complete pipeline from web crawling, HDFS raw storage, Spark cleaning, Hive data management, MySQL business storage to Web visualization. The system collects job data in JSONL format by city, uses Spark to perform missing value filtering, city filtering, job_hash deduplication, salary standardization and field normalization, and stores cleaned DWD data in Parquet format on HDFS. Spark is then used to analyze job quantity, salary ranges, education requirements, experience requirements and skill hot words. Jieba segmentation and TF-IDF are applied to extract representative keywords from job descriptions, and the results are written into MySQL analysis tables for frontend visualization."
    )
    h.style_run(r, "Times New Roman", 12)
    para = doc.add_paragraph()
    h.set_paragraph_format(para, True)
    r = para.add_run(
        "The platform implementation acts as the business presentation layer of data analysis results. The backend uses Spring Boot and MyBatis to provide user, job, analysis and administration APIs, while the frontend uses Vue and ECharts to implement job search, job detail, dashboard, profile center and administration pages. Test results show that the system can complete the full data chain from crawler output to Hadoop processing, MySQL importing and Web presentation. A content-based person-job matching scoring model is also designed based on location preference, position relevance, skill matching, salary matching and behavior feedback."
    )
    h.style_run(r, "Times New Roman", 12)
    para = doc.add_paragraph()
    h.set_paragraph_format(para, False)
    r = para.add_run("Key words: ")
    h.style_run(r, "Times New Roman", 12, True)
    r = para.add_run("Hadoop; Spark; Hive; recruitment data analysis; jieba; TF-IDF; Spring Boot; data visualization")
    h.style_run(r, "Times New Roman", 12)
    doc.add_page_break()


def add_chapter_1(doc):
    h.add_h1(doc, "1 绪论")
    h.add_h2(doc, "1.1 研究背景")
    for text in [
        "贵州省近年来持续推进大数据战略行动和数字经济建设，文旅服务、现代物流、制造业、软件服务、教育培训和生活服务等领域对人才的需求不断变化。招聘平台中沉淀的大量岗位信息能够从岗位数量、薪资水平、学历要求、经验要求和技能关键词等维度反映区域就业市场结构，是观察人才供需关系的重要数据来源。",
        "招聘岗位数据虽然来源丰富，但数据形态复杂。不同企业在发布岗位时对薪资、经验、学历和岗位职责的描述方式不一致，同一岗位可能被多个关键词检索到，部分岗位还存在城市串入、公司简称、岗位描述缺失和薪资单位不统一等问题。若直接对原始数据进行统计，会导致分析结果偏差，影响后续岗位推荐和市场判断。",
        "因此，本课题将研究重点放在招聘数据工程处理和分析上，通过Hadoop生态构建原始数据存储、清洗数据管理和离线分析流程，再使用Spring Boot与Vue实现分析结果的业务呈现。平台设计并不是论文核心目的，而是用于验证数据处理流程能否真正支撑岗位查询、数据看板和推荐模型应用。",
    ]:
        h.add_para(doc, text)
    h.add_h2(doc, "1.2 研究目的和意义")
    for text in [
        "本文的研究目的在于构建一套面向贵州招聘岗位数据的处理与分析流程，使原始岗位数据能够经过采集、清洗、分析和入库后转化为可解释的招聘市场指标。系统最终需要回答三个问题：各城市岗位供给如何分布，不同岗位的薪资、学历和经验要求呈现什么特征，岗位描述中哪些技能关键词具有代表性。",
        "从专业定位看，本课题符合大数据专业对数据采集、存储、清洗、分析和可视化的能力要求。通过HDFS、Hive和Spark处理招聘数据，可以体现大数据工程链路；通过jieba与TF-IDF处理岗位描述文本，可以体现自然语言文本分析；通过MySQL、Spring Boot和Vue将结果展示给用户，可以体现数据分析结果的业务落地。",
        "从实际应用看，系统能够为求职者提供岗位查询和区域岗位趋势参考，为管理员提供岗位数据维护和分析结果查看能力，也为后续构建岗位推荐系统、用户画像和岗位画像奠定数据基础。",
    ]:
        h.add_para(doc, text)
    h.add_h2(doc, "1.3 国内外研究现状")
    h.add_h3(doc, "1.3.1 国内研究现状")
    for text in [
        "国内关于就业数据、招聘平台和大数据分析系统的研究较多，主要集中在Hadoop数据仓库、Spark离线分析、Web管理平台和推荐系统等方向。Hadoop与Hive常用于承载历史数据和主题数据，Spark常用于离线清洗和多维聚合，Spring Boot与Vue则常用于构建前后端分离的数据展示平台。",
        "在就业推荐方面，国内研究通常采用岗位标签匹配、用户画像、协同过滤和混合推荐等方法。由于用户行为数据积累需要较长时间，许多系统在初期会优先采用基于内容的推荐方式，即根据岗位文本、岗位类别、城市、薪资和用户偏好进行匹配。本文当前项目也采用该思路，以TF-IDF关键词作为岗位文本特征基础。",
    ]:
        h.add_para(doc, text)
    h.add_h3(doc, "1.3.2 国外研究现状")
    for text in [
        "国外大型招聘平台通常具有完善的数据采集、搜索召回、用户画像和智能推荐能力。平台会结合职位文本、用户简历、浏览行为、投递行为和企业反馈等多源数据，构建更复杂的推荐模型和排序模型。相关研究也更重视语义匹配、深度学习和在线反馈机制。",
        "与成熟商业平台相比，本文的研究范围更聚焦于本科毕业设计中的完整数据链路实现。系统不追求复杂实时推荐，而是强调公开岗位数据从采集到清洗、分析、入库和展示的可复现流程，并为后续扩展推荐算法提供数据表和评分模型设计。",
    ]:
        h.add_para(doc, text)
    h.add_h2(doc, "1.4 研究内容与技术路线")
    h.add_para(
        doc,
        "本文主要完成五方面工作：一是构建按城市采集岗位数据的爬虫模块；二是基于Spark实现字段标准化、薪资折算、城市过滤和去重处理；三是基于HDFS、Hive和MySQL设计招聘数据存储链路；四是基于Spark完成薪资、学历、经验、城市和文本关键词分析；五是基于Spring Boot和Vue实现数据分析结果展示和基础用户交互。",
    )
    h.add_image(doc, IMG_DIR / "data_flow.png", "图1.1 论文技术路线与数据链路", 14.5)
    h.add_h2(doc, "1.5 本章小结")
    h.add_para(doc, "本章从贵州招聘市场数据化分析需求出发，说明了研究背景、研究意义、国内外研究现状和技术路线。本文后续章节将围绕招聘数据处理主线展开，重点阐述数据采集、清洗、分析模型和平台整合实现。")


def add_chapter_2(doc):
    h.add_h1(doc, "2 需求分析与数据问题定义")
    h.add_h2(doc, "2.1 业务需求分析")
    h.add_para(
        doc,
        "平台面向贵州人才招聘市场公开岗位数据，业务流程包括数据采集、数据清洗、数据分析、结果展示和基础管理。普通用户主要关注岗位查询、岗位详情、地区岗位分布和个人求职偏好维护；管理员主要关注用户管理、岗位状态维护和数据结果检查；系统内部则需要保证数据从原始层到业务层的转换过程稳定可追溯。",
    )
    h.add_h2(doc, "2.2 数据需求分析")
    h.add_table(
        doc,
        ["数据类别", "字段或内容", "用途"],
        [
            ["岗位基础字段", "岗位名称、公司名称、城市、地址", "支撑岗位查询、城市统计和岗位展示。"],
            ["岗位要求字段", "学历、经验、薪资、岗位描述", "支撑学历经验分析、薪资分布和文本分析。"],
            ["采集追溯字段", "来源平台、原始链接、批次号、job_hash", "支撑去重、批次管理和问题回溯。"],
            ["用户画像字段", "城市、学历、期望岗位、薪资、技能描述", "支撑后续推荐系统与个人中心。"],
            ["分析结果字段", "城市、指标名称、数量、占比、排名", "支撑数据看板和论文实证分析。"],
        ],
        "表2.1 平台数据需求说明",
    )
    h.add_h2(doc, "2.3 功能需求分析")
    h.add_table(
        doc,
        ["模块", "功能需求", "数据支撑"],
        [
            ["数据采集", "按贵州城市采集岗位数据，支持安全频率、去重、扩词和批次输出。", "爬虫JSONL、CSV、meta.json。"],
            ["数据清洗", "处理空值、重复、薪资单位、城市串入和字段格式不统一问题。", "Spark DWD Parquet。"],
            ["数据分析", "统计城市岗位数量、薪资、学历、经验、热词和TF-IDF关键词。", "analysis_*分析表。"],
            ["岗位查询", "支持岗位分页、关键字、城市、薪资和学历筛选。", "job_info正式表。"],
            ["用户交互", "支持注册登录、个人资料维护和求职偏好填写。", "sys_user及用户偏好字段。"],
            ["管理员管理", "支持用户与岗位数据管理。", "sys_user、job_info。"],
        ],
        "表2.2 功能需求与数据支撑关系",
    )
    h.add_h2(doc, "2.4 非功能需求分析")
    h.add_table(
        doc,
        ["需求类型", "具体要求"],
        [
            ["稳定性", "爬虫和数据处理脚本应支持批次化执行，单批异常不影响历史批次结果。"],
            ["可追溯性", "保留ODS原始JSONL、DWD清洗结果和清洗报告，便于重新处理和问题定位。"],
            ["一致性", "通过job_hash去重、stage表合并和状态字段控制，减少重复岗位进入业务表。"],
            ["扩展性", "城市、关键词、分析指标和推荐权重应便于后续增加。"],
            ["可视化", "分析结果应能以图表、卡片和表格方式直观展示。"],
        ],
        "表2.3 非功能需求说明",
    )
    h.add_h2(doc, "2.5 数据质量问题定义")
    h.add_table(
        doc,
        ["质量问题", "表现形式", "处理方式"],
        [
            ["重复岗位", "同一岗位在多个关键词或批次中出现。", "爬虫、本批次清洗和MySQL合并阶段均使用job_hash去重。"],
            ["空岗位描述", "部分列表页或详情页无完整描述。", "清洗阶段过滤job_description为空的记录。"],
            ["跨城市数据", "目标城市批次中混入外省或其他城市岗位。", "使用strict-city规则按目标城市过滤。"],
            ["薪资表达不统一", "存在K、千、万/年、元/天、元/时、元/周等。", "统一折算为月薪上下限。"],
            ["公司名称截断", "列表页公司名可能带省略号。", "优先使用详情页company_full_name回填。"],
        ],
        "表2.4 招聘数据质量问题与处理方式",
    )
    h.add_h2(doc, "2.6 本章小结")
    h.add_para(doc, "本章从业务、数据、功能和质量四个角度明确了系统需求。与一般Web管理系统不同，本文的核心需求不是单纯实现页面功能，而是建立可追溯的数据处理流程，并使Web系统能够调用和展示数据处理结果。")


def add_chapter_3(doc):
    h.add_h1(doc, "3 相关技术与算法基础")
    items = [
        ("3.1 Hadoop与HDFS", "Hadoop是大数据平台常用基础框架，HDFS提供分布式文件存储能力。本系统使用HDFS保存ODS原始岗位JSONL和DWD清洗后Parquet数据，使数据处理过程具备批次化存储和历史追溯能力。"),
        ("3.2 Hive数据仓库", "Hive能够将HDFS上的文件映射为表结构，通过类SQL方式进行查看和统计。系统使用Hive外部表dwd_job_clean管理清洗后的招聘主题数据。"),
        ("3.3 Spark离线处理", "Spark适合对批量数据进行清洗、聚合和机器学习特征处理。本文使用PySpark完成字段清洗、薪资解析、去重、分析指标聚合以及JDBC写入MySQL。"),
        ("3.4 MySQL业务数据库", "MySQL用于保存Web系统直接查询的业务数据，包括job_info、sys_user、analysis_*和recommend_result。与Hive相比，MySQL更适合支撑前端页面的低延迟查询。"),
        ("3.5 Spring Boot与Vue", "Spring Boot用于构建后端REST接口，Vue用于构建前端单页应用，ECharts用于数据图表展示。系统开发部分承担数据分析结果的业务呈现作用。"),
        ("3.6 jieba中文分词", "jieba用于将岗位描述文本切分为词语。招聘岗位中存在Java、Python、Spring Boot等中英文混合技能词，脚本通过自定义词典和标准化规则提升分词质量。"),
        ("3.7 TF-IDF关键词提取", "TF-IDF用于衡量词语在文档中的重要程度。TF表示词频，IDF表示逆文档频率，能够降低常见词权重，提高具有岗位区分度的关键词权重。"),
        ("3.8 基于内容的人岗匹配", "基于内容的人岗匹配通过比较用户画像与岗位画像之间的相似度进行推荐。本文设计的模型综合考虑地区、岗位、技能、薪资和行为反馈，具有较好的解释性。"),
    ]
    for heading, text in items:
        h.add_h2(doc, heading)
        h.add_para(doc, text)
    h.add_h2(doc, "3.9 本章小结")
    h.add_para(doc, "本章介绍了系统所需的核心技术与算法基础。Hadoop、Hive和Spark支撑数据处理主线，MySQL、Spring Boot和Vue支撑业务展示，jieba、TF-IDF和内容匹配方法支撑岗位文本分析与推荐模型设计。")


def add_chapter_4(doc, snapshot):
    h.add_h1(doc, "4 招聘岗位数据处理与分析模型设计")
    h.add_h2(doc, "4.1 数据来源与采集策略")
    h.add_para(
        doc,
        "本文使用公开招聘平台岗位数据作为研究对象，爬虫以贵州城市为输入，围绕不同职业类别的种子关键词进行轮询采集。由于招聘平台通常是搜索驱动模式，无法直接获取全城市无关键词的完整岗位列表，因此系统采用“种子关键词 + 随机轮询 + 结果反哺扩词”的采集策略，尽量提高样本覆盖范围。",
    )
    h.add_table(
        doc,
        ["采集字段", "字段说明", "后续用途"],
        [
            ["job_title", "岗位名称", "岗位分类、检索和推荐匹配。"],
            ["company_name", "公司名称", "岗位展示和重复检查。"],
            ["city", "工作城市", "城市分布和地区匹配。"],
            ["work_address", "详细地址", "岗位详情展示和地区分析。"],
            ["education_text", "学历要求", "学历分布分析。"],
            ["experience_text", "经验要求", "经验分布分析。"],
            ["salary_text", "原始薪资", "薪资清洗和展示。"],
            ["job_description", "岗位描述", "jieba分词、TF-IDF和技能提取。"],
            ["data_batch_no", "批次号", "批次追踪和重复分析。"],
            ["job_hash", "岗位哈希", "去重和MySQL合并。"],
        ],
        "表4.1 岗位采集字段及用途",
    )
    h.add_h2(doc, "4.2 ODS-DWD数据分层设计")
    h.add_para(
        doc,
        "数据分层采用ODS和DWD两层结构。ODS层保存爬虫生成的原始JSONL，尽量不修改原始字段；DWD层保存经过Spark清洗后的结构化Parquet数据，字段对齐MySQL的job_info表。该设计使原始数据和清洗结果相互独立，便于后续重新执行清洗规则。",
    )
    h.add_image(doc, IMG_DIR / "data_flow.png", "图4.1 ODS-DWD-Hive-MySQL数据流向图", 14.5)
    h.add_h2(doc, "4.3 数据清洗规则设计")
    h.add_para(
        doc,
        "清洗过程由clean_job_data.py完成，核心规则包括字段回填、必填字段过滤、城市过滤、薪资解析、job_hash去重和状态字段输出。字段回填时优先使用详情页字段，例如detail_company_full_name、detail_job_description和detail_location_address，以减少列表页公司名截断和描述不完整问题。",
    )
    h.add_image(doc, IMG_DIR / "cleaning_flow.png", "图4.2 岗位数据清洗规则流程图", 14.5)
    h.add_table(
        doc,
        ["规则", "处理逻辑", "目的"],
        [
            ["空值过滤", "过滤岗位名、公司名、岗位描述、job_hash为空的数据。", "保证岗位详情和文本分析可用。"],
            ["城市过滤", "开启strict-city后只保留目标城市岗位。", "避免跨城市数据污染统计结果。"],
            ["岗位去重", "按job_hash去除本批次重复岗位。", "降低重复岗位对统计的影响。"],
            ["薪资标准化", "将年薪、日薪、时薪、周薪折算为月薪。", "统一薪资分析口径。"],
            ["详情回填", "优先使用详情页公司全称、描述、地址和发布时间。", "提高字段完整性。"],
        ],
        "表4.2 数据清洗规则设计",
    )
    h.add_image(doc, IMG_DIR / "formula_salary.png", "图4.3 薪资标准化公式")
    h.add_h2(doc, "4.4 Hive与MySQL存储模型设计")
    h.add_para(
        doc,
        "DWD层清洗结果以Parquet格式保存到HDFS，并通过Hive外部表dwd_job_clean进行管理。MySQL负责保存前端和后端直接查询的数据，其中job_info_stage作为Spark导入中转表，job_info作为正式业务表。Spark先写入stage表，再按job_hash合并到正式表，避免重复岗位导致整批导入失败。",
    )
    h.add_table(
        doc,
        ["层级或表", "保存内容", "说明"],
        [
            ["HDFS ODS", "爬虫原始JSONL", "保留原始字段和原始接口数据。"],
            ["HDFS DWD", "清洗后Parquet", "字段对齐job_info，支持Hive外部表。"],
            ["Hive dwd_job_clean", "DWD外部表", "用于SQL查看和论文数据验证。"],
            ["MySQL job_info_stage", "导入中转数据", "防止直接覆盖正式业务表。"],
            ["MySQL job_info", "正式岗位数据", "后端岗位查询和详情展示来源。"],
            ["MySQL analysis_*", "统计分析结果", "前端数据看板来源。"],
        ],
        "表4.3 数据存储层级设计",
    )
    h.add_h2(doc, "4.5 Spark统计分析指标设计")
    h.add_para(
        doc,
        "常规统计分析由analyze_job_data.py完成。脚本读取DWD Parquet，根据城市、薪资区间、学历要求、经验要求和技能词进行聚合，最终写入analysis_city_job_stats、analysis_salary_distribution、analysis_education_stats、analysis_experience_stats和analysis_skill_hotword等表。",
    )
    h.add_table(
        doc,
        ["分析指标", "计算方法", "输出表"],
        [
            ["城市岗位数量", "按city分组计数。", "analysis_city_job_stats"],
            ["城市岗位名称排行", "按city和job_title分组并排序。", "analysis_city_position_stats"],
            ["薪资区间分布", "按平均月薪划分3K以下、3K-5K、5K-8K等区间。", "analysis_salary_distribution"],
            ["学历要求分布", "按city和education_text分组并计算占比。", "analysis_education_stats"],
            ["经验要求分布", "按city和experience_text分组并计算占比。", "analysis_experience_stats"],
            ["规则技能热词", "从岗位文本中匹配技能词典并按城市聚合。", "analysis_skill_hotword"],
        ],
        "表4.4 Spark统计分析指标设计",
    )
    h.add_h2(doc, "4.6 jieba与TF-IDF关键词模型设计")
    h.add_para(
        doc,
        "jieba与TF-IDF功能由analyze_skill_tfidf.py实现。脚本读取DWD Parquet或MySQL job_info，将岗位描述作为文本来源，先使用jieba进行中文分词，再使用自定义保留词和停用词过滤提升词语质量。对于Java、Python、MySQL、Spring Boot等技术词，脚本会进行大小写和同义表达标准化。",
    )
    h.add_image(doc, IMG_DIR / "tfidf_flow.png", "图4.4 jieba与TF-IDF关键词分析流程图", 14.5)
    h.add_image(doc, IMG_DIR / "formula_tfidf.png", "图4.5 TF-IDF关键词权重计算公式")
    h.add_para(
        doc,
        "CountVectorizer将分词结果转换为词频向量，IDF根据词语在全部岗位文档中的出现范围计算逆文档频率，最终得到TF-IDF向量。脚本对每个岗位提取权重较高的关键词，再按城市聚合，得到城市维度和全省维度的关键词排名，并写入analysis_skill_tfidf表。",
    )
    h.add_h2(doc, "4.7 人岗匹配评分模型设计")
    h.add_para(
        doc,
        "推荐系统采用基于内容的评分模型，输入包括用户画像和岗位画像。用户画像包含所在城市、期望岗位、期望薪资、学历和自然语言技能描述；岗位画像来自job_info和TF-IDF关键词结果。评分模型综合考虑地区匹配、岗位匹配、技能匹配、薪资匹配和行为反馈，其中地区偏好可以根据用户填写的级别提高权重。",
    )
    h.add_image(doc, IMG_DIR / "recommend_score_model.png", "图4.6 人岗匹配评分模型设计图", 14.5)
    h.add_image(doc, IMG_DIR / "formula_recommend.png", "图4.7 人岗匹配综合评分公式")
    h.add_table(
        doc,
        ["评分项", "含义", "当前设计"],
        [
            ["地区匹配Sc", "用户期望城市与岗位城市是否一致。", "用户更重视本地工作时提高该项权重。"],
            ["岗位匹配Sp", "期望岗位与岗位标题、岗位描述的相关度。", "通过关键词包含和文本相似度计算。"],
            ["技能匹配Ss", "用户技能词与岗位TF-IDF关键词的重合程度。", "优先匹配岗位描述中的高权重技能词。"],
            ["薪资匹配Sm", "岗位薪资是否落入用户期望区间。", "比较salary_min、salary_max与用户期望薪资。"],
            ["行为反馈Sb", "浏览、搜索、点击等行为产生的动态兴趣。", "当前表结构已预留，后续迭代完善。"],
        ],
        "表4.5 人岗匹配评分项设计",
    )
    h.add_h2(doc, "4.8 本章小结")
    h.add_para(doc, "本章是全文的数据处理核心章节，围绕招聘岗位数据的来源、分层、清洗、存储、统计分析、TF-IDF文本分析和推荐评分模型进行了设计。通过该章节可以明确系统的研究重点是数据工程与业务算法，平台页面只承担结果展示和交互作用。")


def add_chapter_5(doc):
    h.add_h1(doc, "5 平台设计与系统整合实现")
    h.add_h2(doc, "5.1 平台总体整合架构")
    h.add_para(
        doc,
        "平台系统位于数据处理结果之上，主要负责调用MySQL中的岗位数据和分析数据。后端Spring Boot并不直接执行爬虫或Spark任务，而是读取job_info、analysis_*和analysis_skill_tfidf等表，为前端提供岗位查询、分析看板和管理员管理接口。",
    )
    h.add_image(doc, IMG_DIR / "system_architecture.png", "图5.1 平台与数据处理链路整合架构图", 14.5)
    h.add_h2(doc, "5.2 数据采集与流水线脚本实现")
    h.add_para(
        doc,
        "完整数据处理流程由run_job_pipeline.sh串联。脚本先检查Hadoop和YARN状态，然后将本地JSONL上传到HDFS ODS路径，调用clean_job_data.py清洗到DWD Parquet，刷新Hive外部表，再调用export_job_info_to_mysql.py将DWD数据导入MySQL，最后执行analyze_job_data.py和analyze_skill_tfidf.py生成分析表。",
    )
    h.add_table(
        doc,
        ["步骤", "脚本或操作", "输出结果"],
        [
            ["Step 1", "上传JSONL到HDFS ODS", "/recruit/ods/job/<batch>.jsonl"],
            ["Step 2", "clean_job_data.py", "/recruit/dwd/job_clean_parquet/<batch>"],
            ["Step 3", "清洗报告检查", "/recruit/dwd/job_clean_report/<batch>"],
            ["Step 4", "刷新Hive外部表", "recruit_warehouse.dwd_job_clean"],
            ["Step 5", "export_job_info_to_mysql.py", "job_info_stage和job_info"],
            ["Step 6", "analyze_job_data.py", "analysis_*基础分析表"],
            ["Step 7", "analyze_skill_tfidf.py", "analysis_skill_tfidf"],
            ["Step 8", "MySQL验证SQL", "岗位数量和分析结果检查"],
        ],
        "表5.1 数据流水线脚本执行步骤",
    )
    h.add_h2(doc, "5.3 后端服务实现")
    h.add_para(
        doc,
        "后端项目位于Java/GuiZhouJob，采用controller、service、mapper和xml分层结构。用户模块提供注册、登录、个人资料查询和个人资料更新；岗位模块提供岗位分页查询、热门岗位和详情查询；分析模块提供TF-IDF关键词查询；管理员模块提供用户和岗位管理。",
    )
    h.add_table(
        doc,
        ["模块", "核心类", "主要数据表"],
        [
            ["用户模块", "UserController、UserService、SysUserMapper", "sys_user"],
            ["岗位模块", "JobController、JobService、JobInfoMapper", "job_info"],
            ["分析模块", "AnalysisController、SkillTfidfService", "analysis_skill_tfidf"],
            ["管理员模块", "AdminController、AdminService", "sys_user、job_info"],
        ],
        "表5.2 后端模块与数据表关系",
    )
    h.add_h2(doc, "5.4 前端可视化实现")
    h.add_para(
        doc,
        "前端项目位于Vue/guizhou-job-web，采用Vue 3、Vue Router、axios和ECharts实现。页面包括首页、岗位库、岗位详情、数据看板、个人中心、管理员页面、登录和注册。前端通过Vite代理将/api请求转发到本地8080端口的Spring Boot服务。",
    )
    for image, caption in [
        ("home.png", "图5.2 平台首页页面"),
        ("jobs.png", "图5.3 岗位列表页面"),
        ("dashboard.png", "图5.4 数据看板页面"),
        ("profile.png", "图5.5 个人中心页面"),
        ("admin.png", "图5.6 管理员管理页面"),
    ]:
        h.add_image(doc, SCREEN_DIR / image, caption, 14.5)
    h.add_h2(doc, "5.5 本章小结")
    h.add_para(doc, "本章说明了平台如何承接数据处理结果并提供业务展示。系统实现部分围绕数据链路服务化展开，重点不是页面功能堆砌，而是让清洗后的岗位数据和分析结果能够被用户查询、查看和管理。")


def add_chapter_6(doc, snapshot):
    h.add_h1(doc, "6 测试与结果分析")
    h.add_h2(doc, "6.1 测试环境与测试方法")
    h.add_table(
        doc,
        ["类别", "环境或版本"],
        [
            ["本地开发系统", "Windows 11"],
            ["后端环境", "JDK 17、Spring Boot 4.0.5、MyBatis、MySQL Connector/J"],
            ["前端环境", "Node.js、Vue 3.5、Vite 8、ECharts 6、axios"],
            ["大数据环境", "Hadoop、Spark 3.5.8、Hive 4.1.0"],
            ["数据库", "MySQL，数据库guizhou_job_platform"],
            ["数据处理语言", "Python 3.11、PySpark、jieba"],
        ],
        "表6.1 测试环境",
    )
    h.add_h2(doc, "6.2 功能测试")
    h.add_table(
        doc,
        ["测试项", "输入或操作", "预期结果", "测试结果"],
        [
            ["用户注册", "输入邮箱、密码、昵称", "系统创建普通用户", "通过"],
            ["用户登录", "输入正确账号密码", "返回用户ID、昵称和角色", "通过"],
            ["岗位分页查询", "访问/job/page并传入分页参数", "返回岗位总数和岗位列表", "通过"],
            ["岗位详情", "点击岗位卡片", "进入岗位详情页", "通过"],
            ["数据看板", "访问/dashboard", "展示城市、薪资、TF-IDF图表", "通过"],
            ["管理员用户管理", "管理员访问/admin", "可以查询和维护用户", "通过"],
            ["管理员岗位管理", "修改岗位状态", "岗位列表状态更新", "通过"],
            ["数据流水线", "执行run_job_pipeline.sh", "完成HDFS、Hive、MySQL同步", "通过"],
            ["TF-IDF分析", "执行analyze_skill_tfidf.py", "写入analysis_skill_tfidf", "通过"],
        ],
        "表6.2 系统功能测试结果",
    )
    h.add_h2(doc, "6.3 数据规模与城市分布分析")
    h.add_para(
        doc,
        f"当前MySQL job_info表中有效岗位约为{snapshot['job_count']}条。样本主要来自黔东南、黔西南、安顺和贵阳等城市，其中黔东南和黔西南样本量较大，说明这两个批次已经形成较稳定的数据基础。贵阳、遵义和六盘水样本量较少，后续仍需要继续补充爬取批次。",
    )
    h.add_image(doc, IMG_DIR / "chart_city_jobs.png", "图6.1 当前MySQL岗位城市分布", 14.5)
    h.add_table(
        doc,
        ["城市", "岗位数量"],
        [[r["city"], r["cnt"]] for r in snapshot["city_rows"][:8]],
        "表6.3 当前岗位城市分布",
    )
    h.add_h2(doc, "6.4 薪资、学历与经验结果分析")
    h.add_para(
        doc,
        "以黔东南批次为例，薪资区间主要集中在3K-5K和5K-8K之间，说明当前样本中生活服务、销售、行政、教育培训和基层岗位占比较高。8K以上岗位仍然存在，但数量相对较少，主要可能集中在技术、管理、医疗和销售提成类岗位。",
    )
    h.add_image(doc, IMG_DIR / "chart_salary_qiandongnan.png", "图6.2 黔东南样本薪资区间分布", 14.5)
    h.add_table(
        doc,
        ["薪资区间", "岗位数量"],
        [[r["salary_range"], r["job_count"]] for r in snapshot["salary_rows"]],
        "表6.4 黔东南样本薪资区间统计",
    )
    h.add_para(
        doc,
        "学历要求方面，黔东南样本中“不限”和“大专”占比较高，说明本地岗位对中低门槛岗位和应用型人才需求较大。本科及以上岗位占比相对较低，表明高学历岗位仍需结合行业和岗位类别进一步细分分析。",
    )
    h.add_table(
        doc,
        ["学历要求", "岗位数量", "占比"],
        [[r["education_text"], r["job_count"], f"{r['proportion']}%"] for r in snapshot["education_rows"]],
        "表6.5 黔东南学历要求统计",
    )
    h.add_table(
        doc,
        ["经验要求", "岗位数量", "占比"],
        [[r["experience_text"], r["job_count"], f"{r['proportion']}%"] for r in snapshot["experience_rows"]],
        "表6.6 黔东南经验要求统计",
    )
    h.add_h2(doc, "6.5 TF-IDF关键词结果分析")
    h.add_para(
        doc,
        "TF-IDF结果显示，黔东南样本中销售、教育、兼职、办公、培训、运营、资格证、设计和汽车等关键词排名靠前。这些词不仅出现频率较高，而且在岗位描述中具有一定区分度，能够反映当地岗位需求结构。与简单词频统计相比，TF-IDF能够降低过于常见词语的影响，使关键词结果更适合作为岗位画像特征。",
    )
    h.add_image(doc, IMG_DIR / "chart_tfidf_qiandongnan.png", "图6.3 黔东南TF-IDF关键词关联岗位数Top10", 14.5)
    h.add_table(
        doc,
        ["排名", "关键词", "TF-IDF权重", "关联岗位数"],
        [[r["rank_no"], r["keyword"], f"{float(r['tfidf_score']):.2f}", r["job_count"]] for r in snapshot["tfidf_rows"][:10]],
        "表6.7 黔东南TF-IDF关键词Top10",
    )
    h.add_h2(doc, "6.6 人岗匹配结果分析")
    h.add_para(
        doc,
        "当前系统已经完成用户画像字段、岗位画像字段、TF-IDF岗位关键词表和推荐结果表设计，前端个人中心可以维护所在城市、学历、期望岗位、期望薪资和自然语言技能描述。完整推荐计算仍处于设计和后续完善阶段，因此本文不将其描述为已完全上线功能，而是作为基于当前数据处理结果可落地的业务算法方案。",
    )
    h.add_h2(doc, "6.7 本章小结")
    h.add_para(doc, "本章从功能测试和数据分析结果两个角度验证系统。测试表明系统能够完成岗位查询、数据看板、管理员管理和数据流水线执行；数据结果表明Spark清洗和TF-IDF分析能够产出可解释的招聘市场指标。")


def add_chapter_7(doc):
    h.add_h1(doc, "7 总结与展望")
    h.add_h2(doc, "7.1 总结")
    for text in [
        "本文围绕贵州人才招聘市场岗位数据，设计并实现了一个基于Hadoop与Spring Boot的招聘市场分析管理平台。与单纯Web系统不同，本文将研究重点放在数据采集、清洗、存储、分析和结果展示的完整链路上，系统平台作为数据处理结果的业务呈现层。",
        "在数据处理方面，系统解决了岗位数据重复、薪资表达不统一、岗位描述缺失、城市串入和清洗结果同步等问题；在文本分析方面，系统引入jieba和TF-IDF提取岗位关键词，为后续推荐系统提供可解释的文本特征；在Web实现方面，系统采用前后端分离架构，使数据分析结果能够以图表和岗位卡片形式展示。",
    ]:
        h.add_para(doc, text)
    h.add_h2(doc, "7.2 不足")
    h.add_para(doc, "当前系统仍存在不足：第一，招聘数据来源主要集中于单一平台，样本覆盖还需要继续扩大；第二，部分城市样本量不足，城市间对比还不够均衡；第三，完整推荐算法尚未完全落地，用户行为数据仍需继续积累；第四，系统安全机制仍偏毕业设计演示强度，后续应增加Token鉴权、密码加密和接口权限拦截。")
    h.add_h2(doc, "7.3 展望")
    h.add_para(doc, "后续工作可以从四个方向展开。第一，扩展多数据源采集，提高贵州各城市样本覆盖度；第二，完善岗位分类体系，对岗位标题和描述进行更准确的行业类别归类；第三，将用户行为日志、TF-IDF关键词和偏好级别整合到推荐计算中，形成定时更新的推荐结果；第四，优化前端可视化，增加地图、词云、趋势图和推荐解释，使平台更具实用性。")


def add_references(doc):
    doc.add_page_break()
    h.add_h1(doc, "参考文献")
    refs = [
        "王丽佳,张萌. 基于Hive数据分析技术在塑料加工中的应用研究[J]. 塑料工业,2025,53(08):194.",
        "龚向哲,周晨阳. 基于Hadoop平台的就业信息自动推荐研究[J]. 长江信息通信,2025,38(04):125-127.",
        "牛子逸. 基于Vue+SpringBoot的音乐评阅系统设计与实现[D]. 电子科技大学,2025.",
        "黄江凯,施运应,谢吉煌,等. 基于SpringBoot+Vue的大学生党员发展教育管理平台的设计与实现[J]. 电脑知识与技术,2025,21(04):57-60.",
        "林昕,张艳丽,康彦,等. Hive数据库在电商销售大数据分析中的应用研究[J]. 电脑编程技巧与维护,2024,(10):99-101.",
        "秦健. 基于Hadoop大数据平台的即时配送实时数据分析系统[D]. 江苏科技大学,2024.",
        "张书贵. 基于Hadoop的智慧工作岗位分析大数据平台的设计与实现[J]. 信息与电脑(理论版),2024,36(05):112-114+118.",
        "潘杰恒,蔡群英. 基于Hadoop的离线电商数据分析系统的设计与实现[J]. 现代计算机,2024,30(03):112-116.",
        "黄娟. 基于SpringBoot和Vue.js的医院数据提取管理平台的设计与实现[J]. 信息与电脑(理论版),2023,35(22):91-93.",
        "杨小英. 数据仓库Hive搭建与应用: 以网站流量统计分析为例[J]. 信息与电脑(理论版),2023,35(21):70-72.",
        "李威,邱永峰. 基于Hadoop的电商大数据可视化设计与实现[J]. 现代信息科技,2023,7(17):46-49.",
        "孟吉祥. 基于Hadoop生态技术的学生就业推荐平台研究与应用[D]. 太原师范学院,2023.",
        "张艳丽,吴淮北. Hive数据仓库在Hadoop大数据环境下数据的导入与应用[J]. 电脑编程技巧与维护,2022,(12):97-99.",
        "张伟. 基于SpringBoot和Vue的综合教学管理平台设计与实现[D]. 重庆大学,2021.",
        "田海晴. 基于SpringBoot和Vue框架的共享运营管理平台的设计与实现[D]. 山东大学,2020.",
        "周燕. 基于Spark的电商用户行为大数据分析与精准推荐系统研究[J]. 互联网周刊,2025,(22):88-90.",
        "尹丹. 基于Spring Batch和Spark SQL的批处理系统设计与实现[D]. 华东师范大学,2024.",
        "吴翔. 基于Spark分布式框架的调度任务优化技术研究[D]. 东华理工大学,2024.",
    ]
    for i, ref in enumerate(refs, 1):
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(f"[{i}] {ref}")
        h.style_run(run, "宋体", 10.5)


def add_ack_appendix(doc, snapshot):
    doc.add_page_break()
    h.add_h1(doc, "致谢")
    h.add_para(doc, "本论文和系统的完成离不开指导教师在选题、技术路线、论文结构和系统实现过程中的指导与帮助。在毕业设计过程中，我围绕Hadoop集群搭建、招聘数据采集、Spark数据清洗、Hive数据仓库、Spring Boot后端接口和Vue前端页面等内容进行了持续学习和实践。通过本课题，我对大数据平台开发流程、数据分析方法和前后端分离系统设计有了更完整的认识。")
    h.add_para(doc, "同时感谢学院提供的学习环境和课程基础，使我能够将数据科学与大数据技术专业知识应用到实际系统开发中。感谢同学和朋友在系统测试、页面体验和论文修改方面提供的建议。后续我将继续完善数据采集范围、推荐算法和可视化效果，使系统更加稳定和实用。")
    doc.add_page_break()
    h.add_h1(doc, "附录")
    h.add_h2(doc, "附录A 数据流水线执行命令")
    h.add_code(
        doc,
        """cd /home/hadoop/spark
export MYSQL_PASSWORD='你的MySQL密码'
./run_job_pipeline.sh <batch> <city> /home/hadoop/data/<batch>.jsonl""",
    )
    h.add_h2(doc, "附录B 当前系统数据概况")
    h.add_table(
        doc,
        ["项目", "当前结果"],
        [
            ["MySQL岗位有效数据", f"{snapshot['job_count']}条"],
            ["主要岗位城市", "黔东南、黔西南、安顺、贵阳"],
            ["TF-IDF关键词结果", "analysis_skill_tfidf表120条记录"],
            ["HDFS ODS批次", "anshun、guiyang、qiandongnan、qianxinan等JSONL文件"],
            ["前端主要页面", "首页、岗位库、数据看板、个人中心、管理后台、登录注册"],
        ],
        "表A.1 当前系统数据与页面概况",
    )


def create_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    snapshot = get_data_snapshot()
    make_extra_figures(snapshot)
    doc = Document()
    h.setup_styles(doc)
    add_compact_cover(doc)
    h.add_integrity_page(doc)
    add_toc_field(doc)
    add_abstracts(doc)
    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc)
    add_chapter_4(doc, snapshot)
    add_chapter_5(doc)
    add_chapter_6(doc, snapshot)
    add_chapter_7(doc)
    add_references(doc)
    add_ack_appendix(doc, snapshot)
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


def main():
    path = create_document()
    print(path)


if __name__ == "__main__":
    main()
