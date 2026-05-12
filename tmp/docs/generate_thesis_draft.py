from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("D:/graduation Project")
OUT_DIR = Path("D:/毕设/论文初稿")
IMG_DIR = ROOT / "tmp/docs/generated_figures"
SCREEN_DIR = ROOT / "tmp/docs/screens"
OUT_DIR.mkdir(parents=True, exist_ok=True)
IMG_DIR.mkdir(parents=True, exist_ok=True)

TITLE = "基于Hadoop与Spring Boot的贵州人才招聘市场分析管理平台的设计与实现"
AUTHOR = "邰旭东"
STUDENT_NO = "202249020315"
COLLEGE = "大数据学院"
MAJOR = "数据科学与大数据技术"
CLASS_NAME = "大数据222"
ADVISOR = "张德跃"
DATE_TEXT = "2026年4月22日"
OUTPUT = OUT_DIR / f"{TITLE}-初稿.docx"


def chinese_font(size=24, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/msyh.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def draw_box(draw, xy, text, fill, outline, font, text_fill=(25, 55, 46), radius=18):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=2)
    x1, y1, x2, y2 = xy
    lines = text.split("\n")
    line_h = font.size + 8
    total_h = line_h * len(lines)
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        w = bbox[2] - bbox[0]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, fill=text_fill, font=font)
        y += line_h


def draw_arrow(draw, start, end, color=(35, 112, 91), width=4):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if ex >= sx:
        points = [(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)]
    else:
        points = [(ex, ey), (ex + 14, ey - 8), (ex + 14, ey + 8)]
    draw.polygon(points, fill=color)


def make_diagrams():
    font_title = chinese_font(34, True)
    font_box = chinese_font(24, True)
    font_small = chinese_font(20)

    img = Image.new("RGB", (1600, 900), (252, 249, 239))
    draw = ImageDraw.Draw(img)
    draw.text((80, 45), "系统总体架构", fill=(20, 60, 50), font=font_title)
    layers = [
        ("前端展示层\nVue 3 + ECharts\n首页、岗位库、数据看板、个人中心、管理后台", (90, 130, 1510, 245)),
        ("业务服务层\nSpring Boot + MyBatis\n用户、岗位、分析、管理员接口", (90, 275, 1510, 390)),
        ("业务数据库层\nMySQL\njob_info、sys_user、analysis_*、recommend_result", (90, 420, 1510, 535)),
        ("离线分析层\nSpark + jieba + TF-IDF\n薪资、学历、经验、城市、技能关键词分析", (90, 565, 1510, 680)),
        ("大数据存储层\nHDFS + Hive\nODS 原始数据、DWD 清洗数据、主题外部表", (90, 710, 1510, 825)),
    ]
    colors = [(232, 244, 235), (245, 238, 214), (234, 244, 242), (251, 232, 210), (229, 239, 224)]
    for i, (text, box) in enumerate(layers):
        draw_box(draw, box, text, colors[i], (78, 128, 105), font_box)
        if i < len(layers) - 1:
            draw_arrow(draw, ((box[0] + box[2]) // 2, box[3] + 10), ((box[0] + box[2]) // 2, box[3] + 28))
    img.save(IMG_DIR / "system_architecture.png")

    img = Image.new("RGB", (1600, 650), (252, 249, 239))
    draw = ImageDraw.Draw(img)
    draw.text((80, 45), "数据处理流程", fill=(20, 60, 50), font=font_title)
    boxes = [
        ("爬虫采集\nJSONL/CSV", (70, 180, 250, 310)),
        ("HDFS ODS\n原始层", (310, 180, 490, 310)),
        ("Spark清洗\n去重/标准化", (550, 180, 760, 310)),
        ("HDFS DWD\nParquet", (820, 180, 1000, 310)),
        ("Hive外部表\n分析查询", (1060, 180, 1240, 310)),
        ("MySQL\n业务库", (1300, 180, 1480, 310)),
    ]
    for text, box in boxes:
        draw_box(draw, box, text, (239, 246, 237), (78, 128, 105), font_small)
    for i in range(len(boxes) - 1):
        b1 = boxes[i][1]
        b2 = boxes[i + 1][1]
        draw_arrow(draw, (b1[2] + 12, (b1[1] + b1[3]) // 2), (b2[0] - 12, (b2[1] + b2[3]) // 2))
    draw_box(draw, (420, 420, 1180, 540), "分析结果通过 Spring Boot 接口返回前端，前端以图表和岗位卡片形式展示", (255, 248, 226), (198, 141, 83), font_small)
    img.save(IMG_DIR / "data_flow.png")

    img = Image.new("RGB", (1600, 760), (252, 249, 239))
    draw = ImageDraw.Draw(img)
    draw.text((80, 45), "人岗匹配设计流程", fill=(20, 60, 50), font=font_title)
    boxes = [
        ("用户画像\n城市/岗位/薪资/技能", (80, 160, 360, 300)),
        ("岗位画像\n描述/学历/经验/薪资", (80, 430, 360, 570)),
        ("jieba分词\n关键词标准化", (520, 160, 800, 300)),
        ("TF-IDF\n技能权重计算", (520, 430, 800, 570)),
        ("加权评分\n地区+岗位+技能+薪资+行为", (970, 290, 1290, 450)),
        ("推荐结果\n排序/理由/展示", (1370, 290, 1530, 450)),
    ]
    for text, box in boxes:
        draw_box(draw, box, text, (239, 246, 237), (78, 128, 105), font_small)
    draw_arrow(draw, (360, 230), (520, 230))
    draw_arrow(draw, (360, 500), (520, 500))
    draw_arrow(draw, (800, 230), (970, 350))
    draw_arrow(draw, (800, 500), (970, 390))
    draw_arrow(draw, (1290, 370), (1370, 370))
    img.save(IMG_DIR / "recommend_flow.png")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=10.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def style_run(run, font="宋体", size=12, bold=False, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_format(p, first_indent=True, line_spacing=1.5, space_after=0):
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    if first_indent:
        pf.first_line_indent = Pt(24)


def add_para(doc, text, first_indent=True):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_indent=first_indent)
    run = p.add_run(text)
    style_run(run, "宋体", 12)
    return p


def add_center(doc, text, font="宋体", size=12, bold=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    style_run(run, font, size, bold)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    style_run(run, "黑体", 15, True)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, "黑体", 14, True)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 3"]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    style_run(run, "黑体", 12, True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, "黑体", 10.5, False)
    return p


def add_image(doc, path, caption, width_cm=14.5):
    if Path(path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Cm(width_cm))
        add_caption(doc, caption)


def add_table(doc, headers, rows, caption=None, widths=None):
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        style_run(run, "黑体", 10.5)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr_cells[i], "E7EFE8")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), size=10.5)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(code)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    return p


def setup_styles(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    add_center(doc, "贵州理工学院", "宋体", 26, True)
    add_center(doc, "本科毕业论文（设计）", "宋体", 26, True)
    add_center(doc, "（2026届）", "宋体", 18, True)
    for _ in range(2):
        doc.add_paragraph()
    add_center(doc, "设计题目：", "宋体", 16, True)
    add_center(doc, TITLE, "宋体", 18, True)
    for _ in range(2):
        doc.add_paragraph()
    info = [
        f"学    院：{COLLEGE}",
        f"专    业：{MAJOR}",
        f"班    级：{CLASS_NAME}",
        f"学    号：{STUDENT_NO}",
        f"学生姓名：{AUTHOR}",
        f"第一指导教师：{ADVISOR}",
        "第二指导教师：",
    ]
    for item in info:
        p = add_center(doc, item, "宋体", 14)
        p.paragraph_format.space_after = Pt(8)
    for _ in range(2):
        doc.add_paragraph()
    add_center(doc, DATE_TEXT, "宋体", 14)
    doc.add_page_break()


def add_integrity_page(doc):
    add_center(doc, "贵州理工学院本科毕业论文（设计）", "宋体", 16, True)
    add_center(doc, "诚信责任书", "宋体", 18, True, 18)
    add_para(doc, "本人郑重声明：本人所呈交的毕业论文（设计），是在导师的指导下独立进行研究所完成。毕业论文（设计）中凡引用他人已经发表或未发表的成果、数据、观点等，均已明确注明出处。", True)
    add_para(doc, "特此声明。", True)
    for _ in range(5):
        doc.add_paragraph()
    add_para(doc, "论文（设计）作者签名：", False)
    add_para(doc, f"日期：{DATE_TEXT}", False)
    doc.add_page_break()


def add_toc(doc):
    add_center(doc, "目录", "黑体", 15, True, 10)
    toc = [
        "摘要",
        "Abstract",
        "1 绪论",
        "1.1 研究背景与意义",
        "1.2 国内外研究现状",
        "1.3 研究内容与目标",
        "1.4 技术路线",
        "1.5 论文结构安排",
        "2 系统需求分析",
        "3 相关技术与关键方法",
        "4 系统设计",
        "5 系统实现",
        "6 系统测试与结果分析",
        "7 结论与展望",
        "参考文献",
        "致谢",
        "附录",
    ]
    for item in toc:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(item)
        style_run(run, "宋体", 12)
        run = p.add_run(" " + "." * max(8, 46 - len(item)) + " ")
        style_run(run, "宋体", 12)
        run = p.add_run("待更新")
        style_run(run, "宋体", 12)
    add_para(doc, "说明：初稿目录为手工占位目录，定稿时可在 Word 中根据标题样式更新自动目录和页码。", True)
    doc.add_page_break()


def add_abstracts(doc):
    add_center(doc, "摘要", "黑体", 15, True)
    add_para(doc, "随着贵州数字经济、文旅服务、制造业和现代服务业持续发展，区域招聘信息呈现来源分散、更新频繁、岗位描述非结构化和人才供需匹配难度较高等特点。为提高招聘数据利用效率，本文围绕贵州人才招聘市场公开数据，设计并实现一个基于Hadoop与Spring Boot的贵州人才招聘市场分析管理平台。系统以网络爬虫为数据采集入口，按城市和岗位关键词获取岗位名称、公司名称、工作城市、详细地址、学历要求、经验要求、薪资范围和岗位描述等数据；以HDFS作为原始数据与清洗数据的分布式存储基础，使用Spark完成字段标准化、薪资折算、城市过滤、空值过滤和去重处理，并通过Hive组织招聘主题数据；以MySQL保存业务数据和分析结果，使用Spring Boot与MyBatis构建后端接口，使用Vue与ECharts实现岗位检索、数据看板、个人中心和管理员管理等页面。", True)
    add_para(doc, "在数据分析方面，系统围绕城市岗位数量、薪资区间、学历要求、经验要求和热门技能等维度进行统计，并引入jieba分词和TF-IDF方法对岗位描述文本进行关键词提取，为后续人岗匹配提供文本特征支撑。在推荐设计方面，系统结合用户填写的城市偏好、期望岗位、薪资范围、技能描述和行为记录，构建基于内容的综合评分模型，使推荐结果能够体现地区优先级、岗位相关性、技能匹配度和薪资匹配度。当前系统已完成数据采集、清洗、入库、分析、可视化、用户管理和基础岗位查询等核心功能，能够支撑贵州招聘市场数据的展示与初步分析。测试结果表明，系统能够完成从爬虫JSONL到HDFS、Spark、Hive、MySQL和Web展示的完整数据链路，具有一定的实用价值和扩展空间。", True)
    p = doc.add_paragraph()
    set_paragraph_format(p, first_indent=False)
    run = p.add_run("关键词：")
    style_run(run, "黑体", 12, True)
    run = p.add_run("Hadoop；Spring Boot；Spark；Hive；招聘数据分析；jieba；TF-IDF；数据可视化")
    style_run(run, "宋体", 12)
    doc.add_page_break()

    add_center(doc, "Abstract", "Times New Roman", 15, True)
    p = doc.add_paragraph()
    set_paragraph_format(p, True)
    run = p.add_run("With the development of digital economy, tourism services, manufacturing and modern service industries in Guizhou, recruitment information becomes increasingly distributed, frequently updated and highly unstructured. To improve the utilization efficiency of recruitment data, this thesis designs and implements a Guizhou talent recruitment market analysis and management platform based on Hadoop and Spring Boot. The system uses a crawler as the data acquisition module, collects job title, company name, city, work address, education requirement, experience requirement, salary range and job description, stores raw and cleaned data on HDFS, and uses Spark to perform standardization, salary parsing, city filtering, missing value filtering and deduplication. Hive is used to organize the recruitment-themed data warehouse, MySQL is used to store business data and analysis results, Spring Boot and MyBatis are used to provide backend APIs, and Vue with ECharts is used to implement job search, dashboard, profile center and administration pages.")
    style_run(run, "Times New Roman", 12)
    p = doc.add_paragraph()
    set_paragraph_format(p, True)
    run = p.add_run("For data analysis, the system calculates job distribution, salary range, education requirement, experience requirement and hot skill keywords. Jieba segmentation and TF-IDF are introduced to extract keywords from job descriptions, which provide text features for later person-job matching. The recommendation design combines location preference, expected position, salary range, skill description and behavior records to build a content-based weighted scoring model. The current system has implemented data acquisition, cleaning, importing, analysis, visualization, user management and basic job query functions. Test results show that the platform can support a complete data pipeline from crawler JSONL to HDFS, Spark, Hive, MySQL and Web presentation, and has practical value and further extensibility.")
    style_run(run, "Times New Roman", 12)
    p = doc.add_paragraph()
    set_paragraph_format(p, False)
    run = p.add_run("Key words: ")
    style_run(run, "Times New Roman", 12, True)
    run = p.add_run("Hadoop; Spring Boot; Spark; Hive; recruitment data analysis; jieba; TF-IDF; data visualization")
    style_run(run, "Times New Roman", 12)
    doc.add_page_break()


def add_body(doc):
    add_h1(doc, "1 绪论")
    add_h2(doc, "1.1 研究背景与意义")
    add_h3(doc, "1.1.1 研究背景")
    for text in [
        "近年来，贵州省围绕大数据、数字经济、文旅服务、新能源材料、装备制造和现代服务业持续推进产业升级，企业岗位需求、劳动者求职行为和区域人才流动均呈现出更明显的数据化特征。招聘平台上每天产生大量岗位数据，这些数据包含岗位名称、薪资待遇、学历要求、经验要求、岗位描述、工作城市和企业信息等内容，能够从侧面反映区域产业结构、企业用人需求和人才供需变化。",
        "然而，招聘数据具有明显的非结构化和分散化特点。不同平台的字段名称、岗位描述方式、薪资单位和地区表达并不完全一致，同一企业或同一岗位还可能在不同时间和不同关键词下重复出现。若仅依赖人工浏览或简单表格统计，难以形成稳定的数据分析链路，也难以支撑后续推荐系统、岗位画像和人才画像等应用。因此，有必要构建一个面向贵州招聘市场的数据采集、清洗、分析与展示平台。",
        "本课题将Hadoop生态技术与Web应用开发技术结合起来，利用HDFS存储招聘原始数据和清洗数据，利用Hive组织主题数据，利用Spark完成离线清洗与统计分析，利用Spring Boot和Vue实现后端接口与前端展示。通过该平台，可以将爬虫采集的岗位数据转化为可查询、可分析、可展示的数据资源，为贵州人才招聘市场分析和人岗匹配研究提供技术支撑。",
    ]:
        add_para(doc, text)
    add_h3(doc, "1.1.2 研究意义")
    for text in [
        "从理论意义看，本课题将大数据处理、自然语言处理和Web系统设计应用于区域招聘市场分析场景，能够验证Hadoop、Hive、Spark、jieba和TF-IDF等技术在招聘文本分析中的组合使用方式。通过对岗位文本、薪资字段、学历要求和城市分布等数据进行结构化处理，可以为区域人才供需分析提供可复用的工程流程。",
        "从应用意义看，平台能够把招聘岗位数据按照城市、薪资、学历、经验和技能等维度进行可视化展示，使用户更直观地理解贵州不同城市岗位机会的分布情况。对于求职者而言，系统可以辅助其查看本地岗位、筛选期望岗位并记录个人求职偏好；对于管理者而言，系统可以提供岗位数据维护、用户管理和分析结果查看能力。",
        "从工程实践意义看，系统覆盖爬虫、数据清洗、HDFS、Hive、Spark、MySQL、Spring Boot、Vue和ECharts等多个模块，具有较完整的数据处理链路和前后端分离架构，有助于提升大数据系统开发、数据仓库构建、接口设计和可视化实现的综合实践能力。",
    ]:
        add_para(doc, text)

    add_h2(doc, "1.2 国内外研究现状")
    add_h3(doc, "1.2.1 国内研究现状")
    for text in [
        "国内关于招聘数据分析和就业推荐的研究主要集中在三个方向：一是基于Hadoop、Hive和Spark的大数据存储分析技术，二是基于Spring Boot和Vue的Web管理平台开发，三是基于内容或行为的推荐方法。已有研究表明，Hadoop生态在处理大规模历史数据、日志数据和业务数据方面具有较高的可扩展性；Hive能够将HDFS上的数据组织为类SQL表结构，降低数据统计分析门槛；Spark则适合完成离线批处理和多维统计任务。",
        "在应用系统方面，Spring Boot和Vue的前后端分离架构已经广泛用于教育管理、医院数据管理、电商分析和岗位分析等场景。该架构能够提升接口开发效率和页面交互体验，也便于后续进行模块化扩展。对于招聘市场分析平台而言，后端需要稳定提供岗位检索、分析结果查询、用户信息维护和管理员操作接口，前端需要以岗位卡片、筛选表单和图表组件展示数据结果。",
        "在推荐与文本分析方面，国内就业推荐研究常采用岗位标签、用户画像、关键词匹配、协同过滤或混合推荐方法。由于本科毕业设计的系统数据规模和用户行为数据仍处于积累阶段，本文采用以岗位文本和用户偏好为核心的基于内容匹配方法，并使用jieba分词与TF-IDF提取岗位描述关键词，为后续推荐权重计算提供基础。",
    ]:
        add_para(doc, text)
    add_h3(doc, "1.2.2 国外研究现状")
    for text in [
        "国外招聘平台和人才服务系统更强调大规模数据处理、实时搜索、智能推荐和用户画像建设。成熟平台通常结合搜索引擎、分布式计算、机器学习和在线推荐服务，实现岗位召回、个性化排序、行为反馈和实时职位推送。在大规模场景下，系统不仅关注岗位数据本身，还会结合用户简历、浏览行为、投递行为、企业反馈和岗位热度等多源信息进行推荐。",
        "与国外成熟系统相比，本文所实现的平台更聚焦于本科毕业设计范围内的完整工程链路，即从公开岗位数据采集、清洗、离线分析到Web展示和基础匹配设计。系统不追求复杂的在线机器学习模型，而是强调可解释、可演示、可维护的数据处理流程，为后续继续扩展行为推荐和语义匹配模型保留接口和数据表结构。",
    ]:
        add_para(doc, text)

    add_h2(doc, "1.3 研究内容与目标")
    for text in [
        "本文主要研究内容包括：第一，设计面向贵州城市招聘数据的爬虫模块，支持按城市输入目标样本量，利用多类别种子关键词和结果反哺扩词机制采集岗位信息；第二，设计Spark数据清洗流程，对岗位名称、公司名称、城市、学历、经验、薪资和岗位描述等字段进行标准化处理；第三，构建HDFS、Hive和MySQL协同的数据存储结构，使原始数据、清洗数据、分析结果和业务数据各自归位；第四，基于Spark完成薪资、学历、经验、城市和技能关键词分析；第五，构建Spring Boot后端接口和Vue前端页面，实现岗位展示、数据看板、用户中心和管理员模块；第六，设计基于内容的人岗匹配模型，为后续推荐系统提供理论和数据基础。",
        "系统目标是形成一条可运行的数据链路：爬虫生成JSONL文件，上传至HDFS的ODS层；Spark读取ODS数据并清洗到DWD层；Hive对DWD数据进行外部表管理；Spark分析结果和清洗后的岗位数据通过JDBC写入MySQL；Spring Boot从MySQL读取业务数据和统计结果；Vue前端完成可视化展示与交互。",
    ]:
        add_para(doc, text)

    add_h2(doc, "1.4 技术路线")
    add_para(doc, "本文按照软件工程和数据工程相结合的方式展开。首先进行业务需求和功能需求分析，明确平台需要解决的数据采集、清洗、分析、展示和管理问题；其次进行相关技术选型，确定Hadoop、Hive、Spark、MySQL、Spring Boot、Vue、ECharts、jieba和TF-IDF等技术组合；然后完成系统架构设计、数据库设计、数据处理流程设计和接口设计；最后实现各模块功能，并通过功能测试和数据结果分析验证系统可用性。")
    add_image(doc, IMG_DIR / "data_flow.png", "图1.1 技术路线与数据处理流程图", 14.5)
    add_h2(doc, "1.5 论文结构安排")
    add_para(doc, "全文共分为七章。第一章介绍研究背景、意义、国内外研究现状、研究内容和技术路线；第二章进行系统需求分析；第三章介绍相关技术与关键方法；第四章阐述系统总体架构、功能模块、数据流程、数据库和接口设计；第五章说明系统主要模块实现过程；第六章进行系统测试与结果分析；第七章总结全文工作并提出后续展望。")

    add_h1(doc, "2 系统需求分析")
    add_h2(doc, "2.1 业务需求分析")
    for text in [
        "平台面向贵州人才招聘市场公开岗位数据，核心业务目标是将分散的岗位数据转化为可查询、可分析和可展示的数据资产。系统使用者主要包括普通用户和管理员。普通用户关注岗位检索、岗位详情、数据看板和个人偏好维护；管理员关注用户管理、岗位维护和数据状态检查；数据处理模块则负责支撑爬虫数据到业务数据库的转换。",
        "从业务流程看，系统首先通过爬虫采集贵州各城市岗位数据，生成JSONL和CSV等结果文件；随后将JSONL上传至HDFS原始层，使用Spark脚本进行清洗、过滤、去重和薪资解析；清洗结果以Parquet形式保存到HDFS DWD层，并挂载到Hive外部表中；之后通过Spark JDBC将岗位数据和分析结果同步至MySQL；最后由Spring Boot提供接口，Vue前端完成展示。",
    ]:
        add_para(doc, text)
    add_h2(doc, "2.2 功能需求分析")
    add_h3(doc, "2.2.1 招聘数据合规采集需求")
    add_para(doc, "系统需要采集公开招聘岗位信息，字段包括岗位名称、公司名称、城市、详细地址、学历要求、经验要求、原始薪资、薪资上下限、岗位描述、发布时间、数据批次和岗位哈希。采集过程需要控制请求频率，支持手动登录、断点保存、跨批次去重和结果反哺扩词，避免因重复抓取造成无效数据堆积。")
    add_h3(doc, "2.2.2 数据清洗与预处理需求")
    add_para(doc, "系统需要对原始岗位数据进行统一清洗。清洗规则包括删除岗位名称、公司名称、岗位描述或岗位哈希为空的数据；按job_hash去重；将城市、学历、经验字段规范化；将日薪、周薪、年薪等不同薪资表达折算为月薪；过滤明显异常薪资；对非目标城市数据进行严格过滤。")
    add_h3(doc, "2.2.3 招聘数据存储需求")
    add_para(doc, "系统需要使用HDFS保存原始数据和清洗数据，使用Hive管理招聘主题外部表，使用MySQL存储后端业务表和前端展示所需的统计结果。MySQL中需要包含用户表、岗位表、分析表、推荐结果表、用户行为表、用户搜索记录表和ETL批次日志表等。")
    add_h3(doc, "2.2.4 招聘市场分析需求")
    add_para(doc, "系统需要支持城市岗位数量统计、薪资区间分布、学历要求分布、经验要求分布、岗位技能热词和TF-IDF关键词分析。分析结果既要支持全省维度，也要支持按城市维度查看，以便用户比较不同地区招聘结构差异。")
    add_h3(doc, "2.2.5 人岗匹配需求")
    add_para(doc, "系统需要根据用户填写的自然语言技能描述、期望城市、期望岗位和薪资范围进行岗位匹配。用户可以表达不同偏好级别，例如更希望本地工作时，推荐算法应提高地区匹配权重；当本地岗位同时符合期望岗位和技能要求时，应获得更高推荐分。当前初稿阶段已完成推荐表设计、用户画像字段和前端偏好维护，完整推荐计算可在后续迭代中继续完善。")
    add_h3(doc, "2.2.6 数据可视化需求")
    add_para(doc, "前端需要以直观方式展示岗位数据和分析结果，包括首页岗位概览、岗位库筛选、数据看板、城市独立数据分析、薪资分布、学历分布、经验分布、技能关键词排名和岗位卡片展示等。")
    add_h2(doc, "2.3 非功能需求分析")
    add_table(doc, ["需求类型", "具体要求"], [
        ["稳定性", "爬虫和数据处理脚本需要支持进度保存、异常跳过和批次化执行，避免单次异常导致全部结果丢失。"],
        ["可扩展性", "城市、关键词类别、分析指标和推荐权重应便于后续配置和扩展。"],
        ["易用性", "前端页面应具备清晰导航、表单提示、筛选条件和图表展示，便于普通用户和管理员操作。"],
        ["数据一致性", "通过job_hash去重、stage表过渡和状态字段控制，降低重复数据和异常数据进入业务表的风险。"],
        ["安全性", "管理员功能与普通用户功能区分，用户密码采用哈希字段保存，接口返回时不返回密码字段。"],
    ], "表2.1 非功能需求说明")
    add_h2(doc, "2.4 可行性分析")
    add_h3(doc, "2.4.1 技术可行性")
    add_para(doc, "Hadoop、Hive、Spark、Spring Boot、Vue、MySQL、jieba和TF-IDF均为成熟技术，资料充足，能够覆盖本系统的数据采集、离线处理、统计分析、接口服务和前端展示需求。当前项目已在本地和Hadoop集群环境中完成核心链路验证，说明技术方案可行。")
    add_h3(doc, "2.4.2 经济可行性")
    add_para(doc, "系统主要使用开源软件和个人开发环境实现，不依赖商业服务。Hadoop集群基于虚拟机或实验环境搭建，MySQL、Spark、Hive、Vue和Spring Boot均可免费使用，开发成本较低，适合本科毕业设计。")
    add_h3(doc, "2.4.3 操作可行性")
    add_para(doc, "系统操作流程清晰：爬虫通过命令行输入城市和目标数量，数据清洗通过脚本执行，后端和前端通过常规开发命令启动。普通用户通过Web页面操作，不需要理解底层Hadoop命令，具有较好的可操作性。")

    add_h1(doc, "3 相关技术与关键方法")
    techs = [
        ("3.1 Hadoop分布式存储与计算技术", "Hadoop是大数据处理领域常用的分布式基础框架，主要包含HDFS、YARN和MapReduce等组件。本系统使用HDFS保存招聘原始数据和清洗结果，利用其高容错和分布式存储能力支撑批量岗位数据处理；YARN为Spark任务提供资源调度能力。当前实验环境中Hadoop版本为3.5.0，集群节点包括node1、node2和node3。"),
        ("3.2 Hive数据仓库技术", "Hive能够在HDFS之上提供类SQL的数据仓库能力，将清洗后的Parquet数据映射为外部表，便于进行统计查询和结果校验。系统中Hive主要承担主题数据管理作用，保留DWD层数据结构，为Spark分析和后续扩展提供统一数据口径。"),
        ("3.3 Spark离线分析技术", "Spark是一种基于内存计算的大数据处理框架，适合执行批量清洗、统计分析和机器学习特征处理。系统使用PySpark读取JSONL或Parquet数据，完成字段规范化、薪资解析、去重、分析结果聚合和JDBC写入MySQL。当前集群Spark版本为3.5.8。"),
        ("3.4 Spring Boot后端开发技术", "Spring Boot用于快速构建后端服务，减少传统Spring项目配置工作量。本系统使用Spring Boot 4、MyBatis和MySQL实现用户模块、岗位模块、分析模块和管理员模块，通过统一Result结构返回接口数据。"),
        ("3.5 Vue前端开发与可视化技术", "Vue用于构建前端单页应用，ECharts用于图表展示。系统前端包含首页、岗位库、岗位详情、数据看板、个人中心、登录注册和管理后台等页面，使用axios统一调用后端接口。"),
        ("3.6 MySQL数据存储技术", "MySQL负责保存Web系统直接查询的数据，包括岗位主表、用户表、分析结果表和推荐结果表。相比直接从Hive查询，MySQL更适合支撑前端页面的低延迟接口访问。"),
        ("3.7 爬虫与数据采集方法", "系统爬虫基于浏览器自动化和接口数据解析实现，支持按城市采集岗位数据。由于招聘平台无法直接提供全城市无关键词全量数据，系统采用多类别种子关键词轮询、随机顺序采样、标题反哺扩词、跨批次去重和安全速度模式，在控制风险的同时扩大岗位覆盖范围。"),
        ("3.8 jieba分词与TF-IDF关键词提取方法", "jieba用于中文文本分词，将岗位描述切分为可分析的词语；TF-IDF用于计算词语在岗位文本中的重要程度。TF表示词语在文档中出现的频率，IDF表示词语在全量语料中的区分能力，两者相乘可以突出具有代表性的岗位技能词。"),
        ("3.9 基于内容的人岗匹配方法", "基于内容的人岗匹配通过比较用户画像和岗位画像之间的相似度生成推荐结果。本文设计的匹配模型综合考虑地区、岗位名称、技能关键词、薪资范围和行为偏好，并支持用户偏好级别影响权重。该方法解释性较强，适合当前用户行为数据尚未充分积累的阶段。"),
    ]
    for heading, text in techs:
        add_h2(doc, heading)
        add_para(doc, text)

    add_h1(doc, "4 系统设计")
    add_h2(doc, "4.1 系统总体架构设计")
    add_para(doc, "系统采用分层架构设计，自下而上包括大数据存储层、离线分析层、业务数据库层、业务服务层和前端展示层。大数据存储层由HDFS和Hive组成，负责保存原始岗位数据和清洗后的主题数据；离线分析层由Spark脚本组成，负责数据清洗、统计分析、TF-IDF关键词分析和MySQL同步；业务数据库层使用MySQL保存岗位、用户、分析结果和推荐结果；业务服务层使用Spring Boot封装REST API；前端展示层使用Vue和ECharts展示系统功能。")
    add_image(doc, IMG_DIR / "system_architecture.png", "图4.1 系统总体架构图", 14.5)
    add_h2(doc, "4.2 功能模块设计")
    add_table(doc, ["模块", "主要功能"], [
        ["数据采集模块", "按城市采集岗位数据，支持安全速度、详情页补抓、断点保存、去重和扩词。"],
        ["数据清洗模块", "完成字段清洗、薪资解析、目标城市过滤、空描述过滤和job_hash去重。"],
        ["数据分析模块", "统计城市岗位数量、薪资区间、学历、经验、技能热词和TF-IDF关键词。"],
        ["岗位服务模块", "提供岗位分页查询、热门岗位查询和岗位详情查询接口。"],
        ["用户模块", "提供注册、登录、个人资料查询和个人求职偏好维护功能。"],
        ["管理员模块", "提供用户增删改查、岗位维护和状态管理功能。"],
        ["可视化模块", "以图表和卡片方式展示岗位市场分析结果。"],
        ["推荐模块", "设计基于地区、岗位、技能、行为和薪资的综合评分模型。"],
    ], "表4.1 系统功能模块设计")
    add_h2(doc, "4.3 数据处理流程设计")
    add_para(doc, "数据处理流程以批次为单位。爬虫每次运行生成一个城市加时间戳命名的批次目录，目录中包含CSV、JSONL、详情原始JSONL、调试JSON和meta.json。系统推荐使用JSONL作为ODS输入，因为JSONL同时保留列表字段、详情字段和原始字段，更适合作为后续清洗与追溯的数据基础。")
    add_image(doc, IMG_DIR / "data_flow.png", "图4.2 招聘数据处理流程图", 14.5)
    add_h2(doc, "4.4 数据仓库与数据库设计")
    add_h3(doc, "4.4.1 Hive主题数据仓库设计")
    add_para(doc, "Hive外部表主要挂载HDFS DWD层清洗后的Parquet文件，字段与MySQL岗位主表保持一致。ODS层保存爬虫原始JSONL，DWD层保存清洗后的结构化岗位数据，分析结果则由Spark聚合后写入MySQL的analysis_*表。该设计既保留原始数据追溯能力，又保证前端查询效率。")
    add_h3(doc, "4.4.2 MySQL业务数据库设计")
    add_table(doc, ["字段名", "类型", "说明"], [
        ["id", "bigint", "岗位ID，主键自增"],
        ["job_title", "varchar(200)", "岗位名称"],
        ["company_name", "varchar(200)", "公司名称"],
        ["city", "varchar(50)", "工作城市"],
        ["work_address", "varchar(255)", "详细地址"],
        ["education_text", "varchar(50)", "学历要求"],
        ["experience_text", "varchar(50)", "经验要求"],
        ["salary_text", "varchar(100)", "原始薪资描述"],
        ["salary_min", "decimal(10,2)", "最低月薪"],
        ["salary_max", "decimal(10,2)", "最高月薪"],
        ["job_description", "longtext", "岗位描述"],
        ["publish_time", "datetime", "发布时间"],
        ["data_batch_no", "varchar(50)", "数据批次号"],
        ["job_hash", "varchar(64)", "岗位去重哈希"],
        ["status", "tinyint", "岗位状态"],
    ], "表4.2 job_info岗位信息表")
    add_table(doc, ["表名", "作用"], [
        ["sys_user", "保存用户账号、角色、所在城市、学历、期望岗位、期望薪资和自然语言技能描述。"],
        ["job_info_stage", "Spark导入MySQL时的临时过渡表，用于批量写入后再合并到job_info。"],
        ["analysis_city_job_stats", "保存城市岗位数量统计结果。"],
        ["analysis_salary_distribution", "保存城市薪资区间分布结果。"],
        ["analysis_education_stats", "保存学历要求统计结果。"],
        ["analysis_experience_stats", "保存经验要求统计结果。"],
        ["analysis_skill_hotword", "保存技能热词统计结果。"],
        ["analysis_skill_tfidf", "保存jieba与TF-IDF提取的关键词结果。"],
        ["recommend_result", "保存推荐岗位、评分、排名和推荐理由。"],
    ], "表4.3 主要业务表与分析表")
    add_h2(doc, "4.5 接口设计")
    add_table(doc, ["接口", "方法", "功能"], [
        ["/user/register", "POST", "用户注册"],
        ["/user/login", "POST", "用户登录"],
        ["/user/{id}", "GET", "获取用户资料"],
        ["/user/profile", "PUT", "更新个人资料和求职偏好"],
        ["/job/page", "GET", "岗位分页查询和筛选"],
        ["/job/hot", "GET", "热门岗位查询"],
        ["/job/{id}", "GET", "岗位详情查询"],
        ["/analysis/skill-tfidf", "GET", "查询TF-IDF技能关键词"],
        ["/analysis/skill-tfidf/cities", "GET", "查询技能分析城市列表"],
        ["/admin/users", "GET/POST/PUT/DELETE", "管理员用户管理"],
        ["/admin/jobs", "GET/PUT/DELETE", "管理员岗位管理"],
    ], "表4.4 系统主要接口设计")

    add_h1(doc, "5 系统实现")
    add_h2(doc, "5.1 开发环境与运行环境")
    add_table(doc, ["类别", "环境或版本"], [
        ["本地开发系统", "Windows 11"],
        ["后端环境", "JDK 17、Spring Boot 4.0.5、MyBatis、MySQL Connector/J"],
        ["前端环境", "Node.js、Vue 3.5、Vite 8、ECharts 6、axios"],
        ["数据库", "MySQL 8，数据库名guizhou_job_platform"],
        ["大数据环境", "Hadoop 3.5.0、Spark 3.5.8、Hive 4.1.0"],
        ["数据处理语言", "Python 3.11、PySpark、jieba"],
        ["开发工具", "IntelliJ IDEA、PyCharm、DataGrip、Navicat、Xshell"],
    ], "表5.1 开发与运行环境")
    add_para(doc, "截至初稿生成时，MySQL中job_info表约有1.3万条岗位数据，主要覆盖黔东南、黔西南、安顺和贵阳等城市；HDFS的/recruit/ods/job目录保存多个城市批次JSONL文件，/recruit/dwd目录保存清洗后的DWD结果和清洗报告。")
    add_h2(doc, "5.2 招聘数据采集实现")
    add_para(doc, "爬虫脚本位于Crawler/test_1.0/test1.py。脚本运行时由用户输入城市和目标岗位数量，程序根据城市别名映射到招聘平台城市编码，并按多类别关键词轮询采集岗位。由于平台无法直接提供全城市无关键词岗位列表，系统预置技术、销售、行政、财务、教育、医疗、餐饮、制造、物流、法律、房地产、建筑工程等多类关键词，通过随机轮询扩大覆盖范围。")
    add_para(doc, "为降低重复数据，爬虫使用job_hash作为岗位唯一标识，hash由securityId、encryptJobId、岗位名称、公司名称、地址和薪资等字段组合生成。脚本同时维护本次运行去重集合和跨批次缓存文件，避免同一城市重复运行时反复写入相同岗位。")
    add_para(doc, "为提高覆盖率，爬虫实现了结果反哺扩词机制：每当采集到新岗位，程序会对岗位标题进行标准化处理，删除急招、高薪、双休、五险一金等噪声词，拆分标题中的岗位候选词，统计高频且质量较好的候选词，并加入后续关键词池继续搜索。该机制使爬虫不只依赖固定种子词，而能够根据当前城市岗位实际结果动态扩展。")
    add_code(doc, """def build_job_hash(job, work_address, salary_text):
    primary_key = "|".join([
        job.get("securityId", ""),
        job.get("encryptJobId", ""),
        job.get("jobName", ""),
        job.get("brandName", ""),
        work_address,
        salary_text,
    ])
    return hashlib.sha256(primary_key.encode("utf-8")).hexdigest()""")
    add_para(doc, "爬虫每次运行生成城市与时间戳命名的目录，例如anshun_2604211551。目录中包含主CSV文件、主JSONL文件、详情原始JSONL、调试JSON、详情HTML和meta.json。CSV字段与job_info表保持一致，JSONL保留更多原始字段，后续数据清洗以JSONL作为主要输入。")
    add_h2(doc, "5.3 数据清洗与预处理实现")
    add_para(doc, "数据清洗脚本位于DataCleaning/clean_job_data.py。该脚本使用PySpark读取JSONL或CSV数据，优先使用详情页字段回填岗位描述、公司全称、地址、薪资、学历、经验和发布时间等信息。清洗输出字段严格对齐job_info表，包括job_title、company_name、city、work_address、education_text、experience_text、salary_text、salary_min、salary_max、job_description、publish_time、data_batch_no、job_hash和status。")
    add_para(doc, "薪资清洗是招聘数据处理中的重点。系统支持解析8-12K、4-6千、10-20万/年、150-200元/天、30-50元/时、800-1200元/周等表达，并尽量折算为月薪上下限。对于无法可靠折算的薪资保留原始salary_text，同时salary_min和salary_max可以为空；对于明显异常的数据，例如最高薪资低于最低薪资或远超过合理范围的记录，系统会在清洗阶段过滤。")
    add_para(doc, "为了避免空岗位描述进入数据库，清洗脚本对job_title、company_name、job_description和job_hash进行必填校验，缺失任一字段的记录不写入DWD和MySQL。针对城市串入问题，清洗脚本提供strict_city逻辑，当本次批次默认城市为黔西南时，非黔西南城市会被过滤，从而避免跨城市数据影响统计结果。")
    add_code(doc, """valid_required_df = df.filter(
    (col("job_title").isNotNull()) & (trim(col("job_title")) != "")
    & (col("company_name").isNotNull()) & (trim(col("company_name")) != "")
    & (col("job_description").isNotNull()) & (trim(col("job_description")) != "")
    & (col("job_hash").isNotNull()) & (trim(col("job_hash")) != "")
)
city_required_df = valid_required_df.filter(col("city") == lit(target_city))
clean_df = city_required_df.dropDuplicates(["job_hash"])""")
    add_h2(doc, "5.4 Hive数据仓库构建实现")
    add_para(doc, "清洗后的数据以Parquet格式保存到HDFS DWD路径，例如hdfs:///recruit/dwd/job_clean_parquet/anshun_2604211551。Hive外部表通过LOCATION指向对应DWD路径，使系统能够通过SQL查看清洗结果。该方式不移动底层数据，适合保留批次目录并进行多批次管理。")
    add_para(doc, "当前HDFS中已经形成/recruit/ods和/recruit/dwd两级目录。其中ODS层保存爬虫输出的原始JSONL文件，DWD层保存经过Spark清洗后的Parquet数据和清洗报告。与直接把CSV导入MySQL相比，该设计保留了原始数据追溯能力：当后续发现某个字段清洗规则不合理时，可以重新从ODS层读取原始JSONL并再次执行清洗脚本，而不需要重新运行爬虫。")
    add_table(doc, ["HDFS路径", "作用", "当前示例"], [
        ["/recruit/ods/job", "保存爬虫原始JSONL和历史测试CSV", "anshun_2604211551.jsonl、qiandongnan_2604221328.jsonl"],
        ["/recruit/dwd/job_clean_parquet", "保存Spark清洗后的结构化Parquet数据", "按批次目录保存清洗结果"],
        ["/recruit/dwd/job_clean_report", "保存清洗报告，记录原始数、缺失过滤数、去重数和异常数", "每个批次生成part文件"],
        ["Hive外部表dwd_job_clean", "映射DWD层清洗结果，便于SQL查看和验证", "字段与job_info表保持一致"],
        ["MySQL job_info_stage", "Spark JDBC写入的过渡表", "防止直接覆盖正式岗位表"],
        ["MySQL job_info", "后端查询的正式岗位表", "前端岗位库和数据看板的数据来源"],
    ], "表5.4 HDFS与MySQL数据链路说明")
    add_code(doc, """CREATE EXTERNAL TABLE dwd_job_clean (
  job_title STRING,
  company_name STRING,
  city STRING,
  work_address STRING,
  education_text STRING,
  experience_text STRING,
  salary_text STRING,
  salary_min DOUBLE,
  salary_max DOUBLE,
  job_description STRING,
  publish_time STRING,
  data_batch_no STRING,
  job_hash STRING,
  status INT
)
STORED AS PARQUET
LOCATION 'hdfs:///recruit/dwd/job_clean_parquet/批次号';""")
    add_h2(doc, "5.5 Spark离线分析实现")
    add_para(doc, "Spark分析脚本包括analyze_job_data.py和analyze_skill_tfidf.py。前者负责常规统计分析，后者负责基于jieba和TF-IDF的岗位描述关键词分析。常规统计结果写入analysis_city_job_stats、analysis_salary_distribution、analysis_education_stats、analysis_experience_stats和analysis_skill_hotword等表。")
    add_table(doc, ["分析指标", "说明"], [
        ["城市岗位数量", "统计各城市岗位总量，用于展示区域需求差异。"],
        ["薪资区间分布", "将岗位按3K以下、3K-5K、5K-8K、8K-12K、12K-15K、15K以上和未知等区间统计。"],
        ["学历要求分布", "统计不限、初中及以下、高中、中专/中技、大专、本科、硕士、博士等学历要求。"],
        ["经验要求分布", "统计不限、应届生、1年以内、1-3年、3-5年、5-10年等经验要求。"],
        ["技能热词", "基于岗位标题和岗位描述统计销售、管理、培训、招聘、运营等高频词。"],
        ["TF-IDF关键词", "基于jieba分词和Spark ML计算关键词权重，提取更具区分度的岗位技能词。"],
    ], "表5.2 Spark离线分析指标")
    add_h3(doc, "5.5.1 岗位类别分布分析")
    add_para(doc, "当前系统通过关键词类别、岗位标题和城市统计岗位分布。由于招聘平台本身不提供稳定统一的岗位类别字段，系统先以关键词和岗位标题作为近似分类依据，后续可通过岗位标题分类模型进一步归类。")
    add_h3(doc, "5.5.2 薪资水平分析")
    add_para(doc, "薪资分析基于salary_min和salary_max字段。对于已解析薪资的岗位，系统可以计算参考均薪并统计不同城市的薪资区间数量；对于无法解析的薪资，系统保留在未知区间，避免强行计算造成误导。当前数据中安顺和黔东南的5K-8K岗位数量较多，说明中等薪资岗位是样本中的主要组成。")
    add_h3(doc, "5.5.3 学历与经验要求分析")
    add_para(doc, "学历和经验分析通过标准化后的education_text和experience_text字段完成。学历要求主要集中在不限、大专和本科；经验要求中不限、1-3年和3-5年占比较高，说明招聘市场中既有大量基础岗位，也存在一定数量对经验有要求的专业岗位。")
    add_h3(doc, "5.5.4 热门技能词分析")
    add_para(doc, "热门技能词分析从岗位描述中提取常见技能和业务词。当前样本中，安顺城市排名靠前的热词包括销售、管理、培训、招聘、运营、餐饮、采购、设计和美容等，反映服务业、销售管理和本地生活岗位需求较明显。")
    add_h3(doc, "5.5.5 区域人才需求分析")
    add_para(doc, "区域分析主要比较不同城市岗位数量、薪资区间、学历结构和技能需求。当前数据库中黔东南约5391条、黔西南约4819条、安顺约2752条、贵阳约55条，数据仍处于持续采集阶段，后续继续补充贵阳、遵义、铜仁、毕节等城市样本后，区域比较结果会更完整。")
    add_h2(doc, "5.6 基于jieba与TF-IDF的人岗匹配实现")
    add_para(doc, "jieba与TF-IDF功能位于DataCleaning/analyze_skill_tfidf.py。脚本读取DWD Parquet、Hive表或MySQL job_info表，将job_description作为文本来源，使用jieba进行中文分词，并通过自定义保留词和停用词过滤提升关键词质量。对于Java、Python、MySQL、Spring Boot等英文技术词，脚本会进行大小写和同义表达标准化。")
    add_para(doc, "分词结果进入Spark ML的CountVectorizer和IDF模块。CountVectorizer将岗位分词结果转化为词频向量，IDF根据词语在全量岗位中的分布计算逆文档频率，最终得到TF-IDF向量。脚本对每个岗位提取权重最高的若干关键词，再按城市聚合，得到各城市和全省维度的关键词排名，并写入analysis_skill_tfidf表。")
    add_code(doc, """tokenized_df = base_df.withColumn("tokens", tokenize_udf(col("job_description")))
vectorizer = CountVectorizer(inputCol="tokens", outputCol="tf", vocabSize=5000, minDF=2)
cv_model = vectorizer.fit(tokenized_df)
tf_df = cv_model.transform(tokenized_df)
idf_model = IDF(inputCol="tf", outputCol="tfidf").fit(tf_df)
tfidf_df = idf_model.transform(tf_df)""")
    add_para(doc, "在人岗匹配设计中，用户自然语言输入的技能描述同样可以通过jieba切分并转化为关键词集合。系统将岗位关键词与用户技能关键词进行比较，结合地区偏好、岗位名称相似度和薪资匹配度计算综合分。推荐得分可表示为：final_score = w1 * location_score + w2 * position_score + w3 * skill_score + w4 * behavior_score + w5 * salary_score。其中权重可以根据用户填写的偏好级别调整。")
    add_image(doc, IMG_DIR / "recommend_flow.png", "图5.1 基于内容的人岗匹配流程图", 14.5)
    add_h2(doc, "5.7 后端服务实现")
    add_para(doc, "后端项目位于Java/GuiZhouJob，采用典型的controller、service、mapper和xml分层结构。用户模块提供注册、登录、个人资料查询和个人资料更新；岗位模块提供岗位分页查询、热门岗位和详情查询；分析模块提供TF-IDF关键词查询；管理员模块提供用户和岗位管理。")
    add_table(doc, ["模块", "核心类"], [
        ["用户模块", "UserController、UserService、SysUserMapper、SysUserMapper.xml"],
        ["岗位模块", "JobController、JobService、JobInfoMapper、JobInfoMapper.xml"],
        ["分析模块", "AnalysisController、SkillTfidfService、SkillTfidfMapper.xml"],
        ["管理员模块", "AdminController、AdminService、AdminServiceImpl"],
        ["通用模块", "Result、PageResult、BusinessException、GlobalExceptionHandler"],
    ], "表5.3 后端核心模块")
    add_para(doc, "岗位查询接口支持keyword、city、educationText、experienceText、salaryMin、salaryMax和onlyWithSalary等筛选条件。Mapper层统一过滤status=1且岗位描述不为空的数据，保证前端展示的是有效岗位。管理员岗位查询则允许查看和维护不同状态的岗位。")
    add_h2(doc, "5.8 前端可视化实现")
    add_para(doc, "前端项目位于Vue/guizhou-job-web，采用Vue Router进行页面路由管理，axios封装统一请求，ECharts渲染数据看板。页面包含首页、岗位库、岗位详情、数据看板、登录、注册、个人中心和管理后台。前端通过Vite代理将/api请求转发到Spring Boot后端。")
    add_image(doc, SCREEN_DIR / "home.png", "图5.2 系统首页截图", 14.2)
    add_image(doc, SCREEN_DIR / "jobs.png", "图5.3 岗位库页面截图", 14.2)
    add_image(doc, SCREEN_DIR / "dashboard.png", "图5.4 数据看板页面截图", 14.2)
    add_image(doc, SCREEN_DIR / "profile.png", "图5.5 个人中心页面截图", 14.2)
    add_image(doc, SCREEN_DIR / "admin.png", "图5.6 管理后台页面截图", 14.2)

    add_h1(doc, "6 系统测试与结果分析")
    add_h2(doc, "6.1 测试环境与测试方法")
    add_para(doc, "系统测试在本地Windows环境和Hadoop集群环境下进行。本地环境用于验证Spring Boot接口、Vue页面、MySQL数据和前端交互；Hadoop集群用于验证HDFS数据上传、Spark清洗、Hive外部表和Spark JDBC写入MySQL流程。测试方法包括接口测试、页面功能测试、脚本运行测试、数据库结果校验和可视化展示检查。")
    add_h2(doc, "6.2 功能测试")
    add_table(doc, ["测试项", "输入或操作", "预期结果", "测试结果"], [
        ["用户注册", "输入邮箱、密码、昵称", "系统创建普通用户", "通过"],
        ["用户登录", "输入正确账号密码", "返回用户ID、昵称和角色", "通过"],
        ["岗位分页查询", "访问/job/page并传入分页参数", "返回岗位总数和岗位列表", "通过"],
        ["岗位筛选", "按城市、学历、薪资筛选", "列表按条件刷新", "通过"],
        ["数据看板", "访问/dashboard", "展示城市、薪资、学历、经验和技能图表", "通过"],
        ["个人中心", "维护城市、学历、期望岗位、薪资和技能描述", "保存资料并用于岗位预览", "通过"],
        ["管理员用户管理", "管理员进入/admin页面", "查看、新增、编辑和禁用用户", "通过"],
        ["爬虫采集", "输入城市和目标数量", "输出CSV、JSONL、meta和debug目录", "通过"],
        ["Spark清洗", "上传JSONL并执行run_job_pipeline.sh", "生成DWD Parquet并写入MySQL", "通过"],
        ["TF-IDF分析", "执行analyze_skill_tfidf.py", "analysis_skill_tfidf表生成关键词排名", "通过"],
    ], "表6.1 系统功能测试表")
    add_h2(doc, "6.3 数据分析结果分析")
    add_para(doc, "截至当前初稿阶段，MySQL中job_info有效查询数据约13018条，sys_user中已有3个用户，analysis_skill_tfidf表已有120条关键词分析结果。job_info城市分布中，黔东南约5391条，黔西南约4819条，安顺约2752条，贵阳约55条，六盘水和遵义各有少量测试数据。由于爬虫仍在继续补充城市样本，当前数据更适合用于验证系统链路和局部城市分析，后续定稿前需要继续补充全省各城市数据。")
    add_table(doc, ["城市", "岗位数量"], [
        ["黔东南", "5391"],
        ["黔西南", "4819"],
        ["安顺", "2752"],
        ["贵阳", "55"],
        ["六盘水", "1"],
        ["遵义", "1"],
    ], "表6.2 当前job_info城市岗位数量统计")
    add_para(doc, "薪资分布结果显示，安顺样本中5K-8K岗位数量约1722条，3K-5K约990条，8K-12K约722条；黔东南样本中5K-8K约2250条，3K-5K约1568条，8K-12K约876条。可以看出当前样本中5K-8K是较集中的薪资区间，3K-5K岗位也占有较大比例。")
    add_para(doc, "技能热词结果显示，安顺样本中销售、管理、培训、招聘、运营、餐饮、采购、设计、美容等词排名靠前，说明本地样本以销售服务、管理运营和生活服务类岗位为主。TF-IDF结果中，全省维度出现销售、教育、开发、兼职、培训、普通话、软件、资格证、运营等关键词，能够为岗位技能需求分析和后续推荐提供依据。")
    add_para(doc, "从数据质量角度看，当前样本已经能够验证系统主要链路，但仍存在城市样本不均衡问题。黔东南、黔西南和安顺样本量明显高于贵阳、遵义、铜仁等城市，因此当前数据看板中的全省统计结果会受到样本量较大城市的影响。定稿阶段应继续采集贵阳、遵义、毕节、铜仁、黔南和六盘水等城市数据，并在论文结果分析中说明样本来源、采集时间和样本数量。")
    add_para(doc, "从字段质量角度看，清洗脚本已经过滤岗位描述为空的记录，并对城市进行严格过滤，这能够降低无效岗位对前端展示和分析结果的影响。对于薪资字段，系统保留原始salary_text并计算salary_min和salary_max，既便于前端展示招聘平台原始薪资，又便于后端进行薪资区间统计。当遇到周薪、日薪和时薪时，脚本采用统一折算规则转为月薪区间；当薪资表达无法可靠解析时，系统不强制填充上下限，而是将其归入未知区间。")
    add_table(doc, ["质量问题", "处理方式", "意义"], [
        ["重复岗位", "使用job_hash在爬虫、本批次清洗和MySQL合并阶段多次去重", "减少同一岗位重复统计"],
        ["空岗位描述", "清洗阶段过滤job_description为空的数据", "保证详情页和文本分析可用"],
        ["跨城市数据", "根据default-city执行严格城市过滤", "避免黔西南批次出现外省岗位影响统计"],
        ["薪资单位不统一", "对K、千、万/年、元/天、元/时、元/周进行折算", "支持薪资区间分析"],
        ["公司简称", "优先使用详情页公司全称回填", "提高岗位详情展示质量"],
    ], "表6.3 数据质量问题与处理方式")
    add_h2(doc, "6.4 人岗匹配结果分析")
    add_para(doc, "当前系统已经完成用户画像基础字段、岗位画像字段、TF-IDF岗位关键词表和推荐结果表设计。前端个人中心可以保存用户所在城市、学历、期望岗位、期望薪资和自然语言技能描述，并根据当前资料进行岗位筛选预览。完整推荐算法仍需进一步将用户偏好级别、用户行为日志和TF-IDF关键词结果整合为定时或实时推荐任务。")
    add_para(doc, "从设计角度看，基于内容的人岗匹配方法适合本系统当前阶段。原因是岗位数据已经具备标题、描述、城市、薪资、学历和经验等较完整字段，而用户行为数据仍需要后续积累。通过先使用岗位文本和用户技能文本计算匹配度，可以在缺少大量行为样本的情况下实现可解释推荐。后续可在行为数据积累后加入浏览、搜索、收藏和点击记录，形成混合推荐模型。")
    add_h2(doc, "6.5 本章小结")
    add_para(doc, "本章从运行环境、功能测试、数据分析结果和人岗匹配结果四个方面对系统进行了验证。测试表明，系统能够完成主要Web功能和数据处理流程，当前数据分析结果能够反映部分贵州城市岗位样本的薪资、学历、经验和技能需求特征。由于部分功能仍处于迭代阶段，后续需要继续补充数据样本、完善推荐计算和优化前端展示。")

    add_h1(doc, "7 结论与展望")
    add_h2(doc, "7.1 研究工作总结")
    add_para(doc, "本文围绕贵州人才招聘市场分析需求，设计并实现了一个基于Hadoop与Spring Boot的招聘市场分析管理平台。系统完成了从爬虫采集、HDFS存储、Spark清洗、Hive数据仓库、MySQL业务入库、Spring Boot接口服务到Vue前端展示的完整链路。通过该系统，可以对贵州招聘岗位进行分页查询、城市筛选、薪资分析、学历分析、经验分析和技能关键词分析，并提供用户中心和管理后台等基础业务功能。")
    add_para(doc, "在数据处理方面，系统解决了岗位数据重复、薪资表达不统一、岗位描述缺失、城市串入和清洗结果同步等问题；在文本分析方面，系统引入jieba和TF-IDF提取岗位关键词，为后续推荐系统提供可解释的文本特征；在Web实现方面，系统采用前后端分离架构，使数据分析结果能够以图表和岗位卡片形式展示。")
    add_h2(doc, "7.2 系统不足")
    add_para(doc, "当前系统仍存在一些不足。第一，爬虫受目标平台反爬策略和关键词检索机制影响，无法保证全城市全量岗位覆盖，数据样本仍需要持续补充；第二，部分岗位详情页受登录权限限制，岗位描述可能不完整；第三，推荐模块当前主要完成表结构、用户画像和理论设计，尚需进一步落地完整评分任务；第四，前端可视化仍可以继续丰富，例如增加趋势分析、岗位类别聚类和地图可视化。")
    add_h2(doc, "7.3 后续展望")
    add_para(doc, "后续工作可以从四个方向展开。第一，继续扩展数据来源和城市样本，提高贵州全省岗位覆盖度；第二，完善岗位分类体系，对岗位标题和描述进行更准确的行业类别归类；第三，将用户行为日志、TF-IDF关键词和偏好级别整合到推荐计算中，形成可定时更新的推荐结果；第四，优化前端交互和可视化效果，增加地图、词云、趋势图和个性化推荐解释，使平台更具实用性和展示价值。")


def add_references(doc):
    doc.add_page_break()
    add_h1(doc, "参考文献")
    refs = [
        "王丽佳,张萌. 基于Hive数据分析技术在塑料加工中的应用研究[J]. 塑料工业,2025,53(08):194.",
        "龚向哲,周晨阳. 基于Hadoop平台的就业信息自动推荐研究[J]. 长江信息通信,2025,38(04):125-127.",
        "牛子逸. 基于Vue+SpringBoot的音乐评阅系统设计与实现[D]. 电子科技大学,2025.",
        "黄江凯,施运应,谢吉煌,等. 基于SpringBoot+Vue的大学生党员发展教育管理平台的设计与实现[J]. 电脑知识与技术,2025,21(04):57-60.",
        "林昕,张艳丽,康彦,等. Hive数据库在电商销售大数据分析中的应用研究[J]. 电脑编程技巧与维护,2024,(10):99-101.",
        "秦健. 基于Hadoop大数据平台的即时配送实时数据分析系统[D]. 江苏科技大学,2024.",
        "张书贵. 基于Hadoop的智慧工作岗位分析大数据平台的设计与实现[J]. 信息与电脑(理论版),2024,36(05):112-114+118.",
        "潘杰恒,蔡群英. 基于Hadoop的离线电商数据分析系统的设计与实现[J]. 现代计算机,2024,30(03):112-116.",
        "黄娟. 基于SpringBoot和Vue.js的医院数据提取管理平台的设计与实现[J]. 信息与电脑(理论版),2023,35(22):91-93.",
        "杨小英. 数据仓库Hive搭建与应用: 以网站流量统计分析为例[J]. 信息与电脑(理论版),2023,35(21):70-72.",
        "李威,邱永峰. 基于Hadoop的电商大数据可视化设计与实现[J]. 现代信息科技,2023,7(17):46-49.",
        "孟吉祥. 基于Hadoop生态技术的学生就业推荐平台研究与应用[D]. 太原师范学院,2023.",
        "张艳丽,吴淮北. Hive数据仓库在Hadoop大数据环境下数据的导入与应用[J]. 电脑编程技巧与维护,2022,(12):97-99.",
        "张伟. 基于SpringBoot和Vue的综合教学管理平台设计与实现[D]. 重庆大学,2021.",
        "田海晴. 基于SpringBoot和Vue框架的共享运营管理平台的设计与实现[D]. 山东大学,2020.",
        "周燕. 基于Spark的电商用户行为大数据分析与精准推荐系统研究[J]. 互联网周刊,2025,(22):88-90.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"[{i}] {ref}")
        style_run(run, "宋体", 10.5)


def add_ack_appendix(doc):
    doc.add_page_break()
    add_h1(doc, "致谢")
    add_para(doc, "本论文和系统的完成离不开指导教师在选题、技术路线、论文结构和系统实现过程中的指导与帮助。在毕业设计过程中，我围绕Hadoop集群搭建、招聘数据采集、Spark数据清洗、Hive数据仓库、Spring Boot后端接口和Vue前端页面等内容进行了持续学习和实践。通过本课题，我对大数据平台开发流程、前后端分离系统设计和数据分析可视化方法有了更完整的认识。")
    add_para(doc, "同时，感谢学院提供的学习环境和课程基础，使我能够将数据科学与大数据技术专业知识应用到实际系统开发中。感谢同学和身边朋友在系统测试、页面体验和论文修改方面提供的建议。后续我将继续完善数据采集范围、推荐算法和可视化效果，使系统更加稳定和实用。")
    doc.add_page_break()
    add_h1(doc, "附录")
    add_h2(doc, "附录A 数据清洗与入库脚本执行流程")
    add_code(doc, """# 1. 上传爬虫JSONL到HDFS ODS
hdfs dfs -mkdir -p /recruit/ods/job
hdfs dfs -put -f 本地批次.jsonl hdfs:///recruit/ods/job/批次.jsonl

# 2. Spark清洗到HDFS DWD Parquet
spark-submit --master yarn --deploy-mode client \\
  /home/hadoop/spark/clean_job_data.py \\
  --input hdfs:///recruit/ods/job/批次.jsonl \\
  --output hdfs:///recruit/dwd/job_clean_parquet/批次 \\
  --report hdfs:///recruit/dwd/job_clean_report/批次 \\
  --format jsonl --output-format parquet --default-city 城市

# 3. 导出DWD到MySQL并生成分析表
spark-submit --master yarn --deploy-mode client \\
  /home/hadoop/spark/export_job_info_to_mysql.py \\
  --source-type parquet --source hdfs:///recruit/dwd/job_clean_parquet/批次 \\
  --target-table job_info

# 4. 执行常规分析和TF-IDF分析
spark-submit --master yarn /home/hadoop/spark/analyze_job_data.py ...
spark-submit --master yarn /home/hadoop/spark/analyze_skill_tfidf.py ...""")
    add_h2(doc, "附录B 当前系统数据概况")
    add_table(doc, ["项目", "当前结果"], [
        ["MySQL岗位有效查询数据", "约13018条"],
        ["MySQL用户数据", "3条"],
        ["TF-IDF关键词结果", "120条"],
        ["HDFS ODS批次", "anshun、guiyang、qiandongnan、qianxinan等JSONL文件"],
        ["HDFS DWD目录", "/recruit/dwd/job_clean_parquet、/recruit/dwd/job_clean_report"],
        ["前端主要页面", "首页、岗位库、数据看板、个人中心、管理后台、登录注册"],
    ], "表A.1 当前系统数据与页面概况")


def main():
    make_diagrams()
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_integrity_page(doc)
    add_toc(doc)
    add_abstracts(doc)
    add_body(doc)
    add_references(doc)
    add_ack_appendix(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
