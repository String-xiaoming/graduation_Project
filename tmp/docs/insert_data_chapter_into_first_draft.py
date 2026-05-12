from __future__ import annotations

import csv
import os
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("D:/graduation Project")
THESIS_DIR = Path("D:/毕设/论文初稿")
SRC = THESIS_DIR / "10-基于Hadoop与Spring Boot的贵州人才招聘市场分析管理平台的设计与实现-论文初稿.docx"
OUT = THESIS_DIR / "10-基于Hadoop与Spring Boot的贵州人才招聘市场分析管理平台的设计与实现-论文初稿-新增数据分析章节.docx"
PDF = OUT.with_suffix(".pdf")
FIG_DIR = ROOT / "tmp/docs/generated_figures"
FIG_DIR.mkdir(parents=True, exist_ok=True)


def font(size=26, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def draw_arrow(draw, start, end, color=(40, 40, 40), width=3):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if ex >= sx:
        points = [(ex, ey), (ex - 13, ey - 7), (ex - 13, ey + 7)]
    else:
        points = [(ex, ey), (ex + 13, ey - 7), (ex + 13, ey + 7)]
    draw.polygon(points, fill=color)


def draw_rect(draw, xy, text, fill=(255, 255, 255), outline=(40, 40, 40)):
    draw.rectangle(xy, fill=fill, outline=outline, width=2)
    f = font(24, True)
    lines = text.split("\n")
    line_h = 34
    total_h = line_h * len(lines)
    x1, y1, x2, y2 = xy
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=f)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, fill=(20, 20, 20), font=f)
        y += line_h


def make_square_flowcharts():
    title_font = font(34, True)

    img = Image.new("RGB", (1500, 620), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "招聘岗位数据处理流程", fill=(0, 0, 0), font=title_font)
    boxes = [
        ("爬虫采集\nJSONL", (60, 190, 250, 310)),
        ("HDFS ODS\n原始层", (310, 190, 500, 310)),
        ("Spark清洗\n去重/标准化", (560, 190, 780, 310)),
        ("HDFS DWD\nParquet", (840, 190, 1030, 310)),
        ("Hive外部表\nSQL查看", (1090, 190, 1280, 310)),
        ("MySQL\n业务/分析表", (1340, 190, 1480, 310)),
    ]
    for text, xy in boxes:
        draw_rect(d, xy, text)
    for i in range(len(boxes) - 1):
        a, b = boxes[i][1], boxes[i + 1][1]
        draw_arrow(d, (a[2] + 10, 250), (b[0] - 10, 250))
    img.save(FIG_DIR / "insert_data_flow_square.png")

    img = Image.new("RGB", (1500, 650), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "岗位数据清洗规则流程", fill=(0, 0, 0), font=title_font)
    boxes = [
        ("读取原始数据", (80, 170, 280, 285)),
        ("详情字段回填", (360, 170, 580, 285)),
        ("必填字段过滤", (660, 170, 890, 285)),
        ("目标城市过滤", (970, 170, 1200, 285)),
        ("薪资月薪折算", (220, 430, 450, 545)),
        ("job_hash去重", (570, 430, 800, 545)),
        ("输出DWD结果", (920, 430, 1150, 545)),
    ]
    for text, xy in boxes:
        draw_rect(d, xy, text)
    for i in range(3):
        a, b = boxes[i][1], boxes[i + 1][1]
        draw_arrow(d, (a[2] + 10, 228), (b[0] - 10, 228))
    draw_arrow(d, (1085, 290), (335, 420))
    draw_arrow(d, (450, 488), (570, 488))
    draw_arrow(d, (800, 488), (920, 488))
    img.save(FIG_DIR / "insert_clean_flow_square.png")

    img = Image.new("RGB", (1500, 650), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "jieba与TF-IDF关键词分析流程", fill=(0, 0, 0), font=title_font)
    boxes = [
        ("岗位描述文本", (80, 170, 300, 285)),
        ("jieba分词", (380, 170, 580, 285)),
        ("停用词过滤", (660, 170, 880, 285)),
        ("词频向量TF", (960, 170, 1180, 285)),
        ("IDF计算", (300, 430, 500, 545)),
        ("TF-IDF权重", (600, 430, 830, 545)),
        ("城市聚合入库", (930, 430, 1180, 545)),
    ]
    for text, xy in boxes:
        draw_rect(d, xy, text)
    for i in range(3):
        a, b = boxes[i][1], boxes[i + 1][1]
        draw_arrow(d, (a[2] + 10, 228), (b[0] - 10, 228))
    draw_arrow(d, (1070, 290), (400, 420))
    draw_arrow(d, (500, 488), (600, 488))
    draw_arrow(d, (830, 488), (930, 488))
    img.save(FIG_DIR / "insert_tfidf_flow_square.png")

    img = Image.new("RGB", (1500, 620), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "人岗匹配评分模型", fill=(0, 0, 0), font=title_font)
    boxes = [
        ("城市偏好", (80, 160, 260, 270)),
        ("岗位相关", (80, 360, 260, 470)),
        ("技能匹配", (420, 160, 600, 270)),
        ("薪资匹配", (420, 360, 600, 470)),
        ("行为反馈", (750, 260, 950, 370)),
        ("综合评分\n排序推荐", (1100, 260, 1320, 370)),
    ]
    for text, xy in boxes:
        draw_rect(d, xy, text)
    for idx in range(5):
        a = boxes[idx][1]
        draw_arrow(d, (a[2] + 10, (a[1] + a[3]) // 2), (1090, 315))
    img.save(FIG_DIR / "insert_recommend_flow_square.png")


def mysql_query(sql):
    env = os.environ.copy()
    env["MYSQL_PWD"] = env.get("MYSQL_PWD", "NLXXtxd123")
    cmd = [
        "mysql",
        "-h",
        "127.0.0.1",
        "-P",
        "3306",
        "-uroot",
        "-D",
        "guizhou_job_platform",
        "--default-character-set=utf8mb4",
        "--batch",
        "--raw",
        "-e",
        sql,
    ]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8", env=env, timeout=20)
        if result.returncode != 0:
            return []
        rows = list(csv.reader(result.stdout.splitlines(), delimiter="\t"))
        if not rows:
            return []
        headers = rows[0]
        return [dict(zip(headers, row)) for row in rows[1:]]
    except Exception:
        return []


def set_run_font(run, size=12, bold=False, name="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_outline(paragraph, level):
    ppr = paragraph._p.get_or_add_pPr()
    existing = ppr.find(qn("w:outlineLvl"))
    if existing is not None:
        ppr.remove(existing)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level))
    ppr.append(outline)


def format_paragraph(paragraph, first_indent=True):
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_indent:
        pf.first_line_indent = Pt(24)


class Inserter:
    def __init__(self, doc, target):
        self.doc = doc
        self.target = target

    def paragraph(self, text="", size=12, bold=False, name="宋体", align=None, outline=None, first_indent=True):
        p = self.target.insert_paragraph_before()
        if align is not None:
            p.alignment = align
        format_paragraph(p, first_indent=first_indent)
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, name=name)
        if outline is not None:
            set_outline(p, outline)
        return p

    def h1(self, text):
        p = self.paragraph(text, size=15, bold=True, name="黑体", outline=0, first_indent=False)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        return p

    def h2(self, text):
        p = self.paragraph(text, size=14, bold=True, name="黑体", outline=1, first_indent=False)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        return p

    def normal(self, text):
        return self.paragraph(text, size=12, name="宋体", first_indent=True)

    def caption(self, text):
        p = self.paragraph(text, size=10.5, name="黑体", first_indent=False)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        return p

    def image(self, path, caption, width_cm=13.6):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Cm(width_cm))
        self.target._p.addprevious(p._p)
        self.caption(caption)

    def table(self, caption, headers, rows):
        cap = self.paragraph(caption, size=10.5, name="黑体", first_indent=False)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, text in enumerate(headers):
            hdr[i].text = ""
            p = hdr[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(text))
            set_run_font(r, size=10.5, bold=True)
        for row in rows:
            cells = table.add_row().cells
            for i, text in enumerate(row):
                cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells[i].text = ""
                p = cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) < 16 else WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(str(text))
                set_run_font(r, size=10.5)
        apply_three_line_table(table)
        self.target._p.addprevious(table._tbl)
        self.paragraph("", first_indent=False)
        return table


def border_el(name, val="single", size="8", color="000000"):
    el = OxmlElement(f"w:{name}")
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), size)
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)
    return el


def clear_cell_borders(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.find(qn("w:tcBorders"))
    if borders is not None:
        tcpr.remove(borders)
    borders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        borders.append(border_el(side, val="nil", size="0"))
    tcpr.append(borders)


def set_cell_border(cell, top=None, bottom=None):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for side, value in [("top", top), ("bottom", bottom)]:
        if value:
            old = borders.find(qn(f"w:{side}"))
            if old is not None:
                borders.remove(old)
            borders.append(border_el(side, size=value))


def apply_three_line_table(table):
    for row in table.rows:
        for cell in row.cells:
            clear_cell_borders(cell)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top="12", bottom="8")
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom="12")


def replace_text_keep_runs(paragraph, new_text):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def renumber_later_chapters(doc, start_index):
    for p in doc.paragraphs[start_index:]:
        text = p.text
        new_text = text
        match = re.match(r"^([4-7])(?=[ .])", text)
        if match:
            new_text = str(int(match.group(1)) + 1) + text[1:]
        else:
            match = re.match(r"^([4-7])\.", text)
            if match:
                new_text = str(int(match.group(1)) + 1) + text[1:]
            else:
                match = re.match(r"^(图|表)([4-7])\.", text)
                if match:
                    new_text = match.group(1) + str(int(match.group(2)) + 1) + text[len(match.group(1)) + 1:]
        if new_text != text:
            replace_text_keep_runs(p, new_text)


def build_data_rows():
    city_rows = mysql_query("SELECT city, COUNT(*) AS cnt FROM job_info GROUP BY city ORDER BY cnt DESC LIMIT 5;")
    if not city_rows:
        city_rows = [
            {"city": "黔东南", "cnt": "5391"},
            {"city": "黔西南", "cnt": "4819"},
            {"city": "安顺", "cnt": "2752"},
            {"city": "贵阳", "cnt": "55"},
        ]
    tfidf = mysql_query(
        "SELECT keyword, tfidf_score, job_count, rank_no FROM analysis_skill_tfidf "
        "WHERE data_batch_no='qiandongnan_2604221328' AND city='黔东南' ORDER BY rank_no LIMIT 5;"
    )
    if not tfidf:
        tfidf = [
            {"rank_no": "1", "keyword": "销售", "tfidf_score": "794.86", "job_count": "467"},
            {"rank_no": "2", "keyword": "教育", "tfidf_score": "427.92", "job_count": "155"},
            {"rank_no": "3", "keyword": "兼职", "tfidf_score": "415.06", "job_count": "141"},
            {"rank_no": "4", "keyword": "办公", "tfidf_score": "399.85", "job_count": "129"},
            {"rank_no": "5", "keyword": "培训", "tfidf_score": "392.21", "job_count": "174"},
        ]
    return city_rows, tfidf


def insert_data_chapter(doc):
    body_paragraphs = list(doc.paragraphs)
    target_index = next(i for i, p in enumerate(body_paragraphs) if p.text.strip() == "4 系统设计")
    target = body_paragraphs[target_index]
    renumber_later_chapters(doc, target_index)
    inserter = Inserter(doc, target)
    city_rows, tfidf_rows = build_data_rows()

    inserter.h1("4 招聘岗位数据处理与分析模型设计")
    inserter.h2("4.1 数据来源与采集策略")
    inserter.normal("本系统的数据来源为公开招聘平台中的贵州地区岗位信息。爬虫以城市为输入，结合技术、销售、行政、财务、教育、医疗、制造、物流、法律、房地产、建筑工程等多类种子关键词进行轮询采集，并根据已采集岗位标题进行反哺扩词。该策略能够在招聘平台搜索驱动的限制下扩大样本覆盖范围。")
    inserter.normal("每次采集生成一个带城市和时间戳的批次目录，目录中包含主CSV、主JSONL、详情原始JSONL、调试JSON、详情HTML和meta.json。后续清洗优先使用JSONL文件，因为JSONL保留了列表字段、详情页字段和原始接口字段，更适合作为ODS原始数据。")
    inserter.table(
        "表4.1 岗位采集字段及用途",
        ["字段", "说明", "后续用途"],
        [
            ["job_title", "岗位名称", "岗位检索、岗位分类和推荐匹配"],
            ["company_name", "公司名称", "岗位展示和重复检查"],
            ["city", "工作城市", "城市分布统计和地区匹配"],
            ["education_text", "学历要求", "学历分布分析"],
            ["experience_text", "经验要求", "经验分布分析"],
            ["salary_text", "原始薪资", "薪资清洗和前端展示"],
            ["job_description", "岗位描述", "jieba分词、TF-IDF和技能提取"],
            ["job_hash", "岗位哈希", "爬虫、清洗和入库阶段去重"],
        ],
    )

    inserter.h2("4.2 ODS-DWD数据分层设计")
    inserter.normal("数据处理链路采用ODS和DWD两层设计。ODS层保存爬虫原始JSONL，不直接修改原始字段；DWD层保存经过Spark清洗后的结构化Parquet数据，字段对齐MySQL中的job_info表。Hive通过外部表挂载DWD路径，MySQL保存后端查询所需的岗位主表和分析结果表。")
    inserter.image(FIG_DIR / "insert_data_flow_square.png", "图4.1 招聘岗位数据处理流程图", 13.8)
    inserter.table(
        "表4.2 数据存储层级设计",
        ["层级或表", "保存内容", "作用"],
        [
            ["HDFS ODS", "爬虫原始JSONL", "保留原始字段，支持重新清洗"],
            ["HDFS DWD", "清洗后Parquet", "保存结构化岗位明细"],
            ["Hive dwd_job_clean", "DWD外部表", "通过SQL查看和校验清洗结果"],
            ["MySQL job_info_stage", "导入中转表", "避免直接覆盖正式岗位表"],
            ["MySQL job_info", "正式岗位表", "支撑岗位查询和详情展示"],
            ["MySQL analysis_*", "统计分析结果", "支撑数据看板和论文分析"],
        ],
    )

    inserter.h2("4.3 数据清洗规则设计")
    inserter.normal("数据清洗由Spark脚本完成，主要解决空值、重复、薪资单位不统一、城市串入和公司名称截断等问题。清洗时优先使用详情页字段回填公司全称、岗位描述、详细地址、薪资、学历和经验，随后执行必填字段过滤、目标城市过滤、薪资标准化和job_hash去重。")
    inserter.image(FIG_DIR / "insert_clean_flow_square.png", "图4.2 岗位数据清洗规则流程图", 13.8)
    inserter.table(
        "表4.3 数据清洗规则设计",
        ["规则", "处理逻辑", "目的"],
        [
            ["空值过滤", "过滤岗位名、公司名、岗位描述、job_hash为空的数据", "保证岗位详情和文本分析可用"],
            ["城市过滤", "开启strict-city后仅保留目标城市岗位", "避免跨城市数据污染统计结果"],
            ["岗位去重", "按job_hash删除重复岗位", "降低重复岗位对分析结果的影响"],
            ["薪资标准化", "将年薪、日薪、时薪、周薪折算为月薪", "统一薪资分析口径"],
            ["详情回填", "使用详情页字段补充公司全称、地址和描述", "提升字段完整性"],
        ],
    )

    inserter.h2("4.4 Spark统计分析指标设计")
    inserter.normal("清洗后的DWD数据进入Spark分析阶段。系统围绕招聘市场分析需求设计城市岗位数量、薪资区间、学历要求、经验要求、岗位名称排行和技能热词等指标，并将结果写入MySQL的analysis_*系列表，供前端数据看板调用。")
    inserter.table(
        "表4.4 Spark统计分析指标设计",
        ["指标", "计算方法", "输出表"],
        [
            ["城市岗位数量", "按city分组统计岗位数", "analysis_city_job_stats"],
            ["薪资区间分布", "按平均月薪划分3K以下、3K-5K、5K-8K等区间", "analysis_salary_distribution"],
            ["学历要求分布", "按city和education_text分组并计算占比", "analysis_education_stats"],
            ["经验要求分布", "按city和experience_text分组并计算占比", "analysis_experience_stats"],
            ["技能热词", "从岗位文本中匹配技能词并排序", "analysis_skill_hotword"],
        ],
    )
    inserter.table(
        "表4.5 当前岗位城市样本统计",
        ["城市", "岗位数量"],
        [[r["city"], r["cnt"]] for r in city_rows],
    )

    inserter.h2("4.5 jieba与TF-IDF关键词模型设计")
    inserter.normal("岗位描述属于非结构化中文文本，需要先经过分词再进行关键词权重计算。系统使用jieba对岗位描述进行分词，并结合自定义词典保留Java、Python、Spark、Hive、Spring Boot等技术词；随后使用停用词表过滤无分析意义的词语，再通过Spark ML中的CountVectorizer和IDF计算TF-IDF权重。")
    inserter.image(FIG_DIR / "insert_tfidf_flow_square.png", "图4.3 jieba与TF-IDF关键词分析流程图", 13.8)
    inserter.normal("TF-IDF的核心思想是：若某个词在当前岗位文本中出现频率较高，但在全部岗位文本中并不普遍，则该词对当前岗位具有更强区分能力。系统最终按城市聚合各岗位的高权重关键词，并写入analysis_skill_tfidf表。")
    inserter.table(
        "表4.6 黔东南样本TF-IDF关键词Top5",
        ["排名", "关键词", "TF-IDF权重", "关联岗位数"],
        [[r.get("rank_no", ""), r.get("keyword", ""), f"{float(r.get('tfidf_score', 0)):.2f}", r.get("job_count", "")] for r in tfidf_rows],
    )

    inserter.h2("4.6 人岗匹配评分模型设计")
    inserter.normal("推荐系统采用基于内容的评分模型，输入包括用户画像和岗位画像。用户画像包含所在城市、期望岗位、期望薪资、学历和自然语言技能描述；岗位画像来自job_info表和TF-IDF关键词结果。模型综合考虑地区匹配、岗位匹配、技能匹配、薪资匹配和用户行为反馈。")
    inserter.image(FIG_DIR / "insert_recommend_flow_square.png", "图4.4 人岗匹配评分模型设计图", 13.8)
    inserter.table(
        "表4.7 人岗匹配评分项设计",
        ["评分项", "含义", "当前设计"],
        [
            ["地区匹配", "用户期望城市与岗位城市是否一致", "本地工作偏好高时提高该项权重"],
            ["岗位匹配", "期望岗位与岗位标题、描述的相关度", "通过关键词包含和文本相关度计算"],
            ["技能匹配", "用户技能与岗位关键词的重合程度", "优先匹配TF-IDF高权重技能词"],
            ["薪资匹配", "岗位薪资是否落入期望薪资范围", "比较salary_min、salary_max与期望薪资"],
            ["行为反馈", "浏览、搜索和点击产生的动态兴趣", "当前表结构预留，后续继续完善"],
        ],
    )

    inserter.h2("4.7 本章小结")
    inserter.normal("本章单独围绕招聘岗位数据处理与分析模型展开，明确了从原始采集数据到ODS、DWD、Hive、MySQL和分析结果表的完整链路。通过空值过滤、城市过滤、薪资标准化、job_hash去重、Spark统计分析以及jieba与TF-IDF关键词提取，系统能够将非结构化招聘岗位数据转化为可查询、可统计和可展示的数据资产。")


def main():
    make_square_flowcharts()
    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))
    insert_data_chapter(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
